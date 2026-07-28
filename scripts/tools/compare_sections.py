from docx import Document

b = Document(r"E:\OC DOCS\Comprendre_Optimum_Control_backup.docx")
v = Document(r"E:\OC DOCS\Comprendre_Optimum_Control_V2.docx")

def dump(doc, label, needles):
    lines = [f"=== {label} ==="]
    for needle in needles:
        for i, p in enumerate(doc.paragraphs):
            if needle in p.text:
                lines.append(f"\n[{i}] {p.text.strip()[:70]}")
                for j in range(i+1, min(i+8, len(doc.paragraphs))):
                    t = doc.paragraphs[j].text.strip()
                    s = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
                    if not t: t = "(vide)"
                    lines.append(f"  {j}|{s}|{t[:65]}")
                break
    return lines

needles = [
    "Détailler les pertes",
    "Vocabulaire essentiel",
    "Raccourcis clavier",
    "Les quatre états",
    "Méthodes de tri",
    "Volume et poids",
    "Les achats de Joe",
    "Calcul du coût unitaire",
]
open(r"E:\OC DOCS\compare_sections.txt", "w", encoding="utf-8").write(
    "\n".join(dump(b, "BACKUP", needles) + ["\n"] + dump(v, "V2", needles))
)
