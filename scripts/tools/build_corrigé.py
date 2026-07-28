#!/usr/bin/env python3
"""Génère le corrigé détaillé de l'exercice Comptoir du Marché."""
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"E:\OC DOCS\Exercice_Corrige_Detaille.docx")
CALC = Path(r"E:\OC DOCS\corrigé_calc.txt")

def add_title(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_p(doc, text):
    doc.add_paragraph(text)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

def build():
    subprocess.run(["python", r"E:\OC DOCS\calc_corrigé.py"], check=True)
    calc_text = CALC.read_text(encoding="utf-8") if CALC.exists() else ""

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CORRIGÉ DÉTAILLÉ\n")
    r.bold = True
    r.font.size = Pt(20)
    t.add_run("Exercice pratique — Le Comptoir du Marché\n\n").font.size = Pt(14)
    t.add_run("Ne consultez ce document qu'après avoir tenté l'exercice.\n").font.size = Pt(11)
    doc.add_page_break()

    add_title(doc, "1. Coûts unitaires des items")
    add_p(doc, "Formule : coût caisse ÷ quantité en unité recette × (100 ÷ yield %).")
    add_table(doc, ["Item", "Calcul", "Coût unitaire"], [
        ["Pain burger", "2,40 $ ÷ 12", "0,20 $ / chacun"],
        ["Bœuf 80/20", "42,00 $ ÷ 5000 g", "0,0084 $ / g"],
        ["Bacon", "18,50 $ ÷ ~66 tranches (2×500 g)", "~0,295 $ / tranche"],
        ["Mayonnaise", "48,00 $ ÷ 15 120 ml", "0,00318 $ / ml"],
        ["Ketchup", "38,00 $ ÷ ~20 232 ml", "0,00188 $ / ml"],
        ["Pommes de terre", "22,00 $ ÷ (50 lb en g), yield 85 %", "0,00114 $ / g utilisable"],
        ["Huile friture", "38,00 $ ÷ 16 000 ml", "0,00238 $ / ml"],
        ["Boîte burger", "45,00 $ ÷ 500", "0,09 $ / chacun"],
        ["Serviette", "35,00 $ ÷ 5000", "0,007 $ / chacun"],
        ["Cola", "12,00 $ ÷ 24", "0,50 $ / canette"],
    ])

    add_title(doc, "2. Coûts des preps")
    add_table(doc, ["Prep", "Calcul", "Coût unité sortie"], [
        ["Boulette 80 g", "77 g bœuf + 3 g sel", "≈ 0,652 $ / boulette"],
        ["Frites maison — batch", "(5000 g pommes + 200 ml huile + 50 g sel) ÷ 4000 g rendu", "~0,00157 $ / g frite"],
        ["Portion frites 200 g", "200 g × coût/g batch", "~0,32 $ / portion"],
        ["Sauce burger", "(800 ml mayo + 200 ml ketchup) ÷ 950 ml", "~0,00307 $ / ml"],
        ["Chili — batch", "Haricots + bœuf chili + tomates + oignons + sel → 12 000 g", "~0,00309 $ / g"],
    ])

    add_title(doc, "3. Coûts recette et food cost %")
    add_p(doc, "Valeurs calculées (arrondis OC peuvent différer légèrement) :")
    add_table(doc, ["Product", "Coût recette", "Prix vente", "Food cost %", "Commentaire"], [
        ["Burger classique", "~1,13 $", "14,95 $", "~7,6 %", "Nourriture seule ; papier en sus"],
        ["Burger bacon", "~1,72 $", "16,95 $", "~10,1 %", "+ 2 tranches bacon"],
        ["Frites moyennes", "~0,32 $", "5,50 $", "~5,8 %", "1 Portion frites 200 g (prep) + serviette"],
        ["Salade César", "~1,43 $", "13,50 $", "~10,6 %", "4 items — laitue, dressing, parmesan, croûtons"],
        ["Bol chili", "~1,25 $", "9,95 $", "~12,5 %", "400 g chili + fourchette"],
        ["Cola", "0,50 $", "3,25 $", "15,4 %", "Item direct"],
        ["Eau", "0,33 $", "2,75 $", "12,1 %", "Item direct"],
    ])
    add_p(doc, "Note pédagogique : avec les prix d'achat de l'exercice, le food cost est nettement sous la cible de 32 % — vos marges sont confortables. Si OC affiche un % beaucoup plus élevé, vérifiez les unités (grammes vs kilos, ml vs litres) ou un yield oublié.")

    add_title(doc, "4. Détail — Burger classique")
    add_table(doc, ["Ingrédient", "Qté", "Coût unit.", "Coût ligne"], [
        ["Pain", "1", "0,20 $", "0,20 $"],
        ["Boulette (prep)", "1", "≈ 0,652 $", "≈ 0,652 $"],
        ["Sauce maison", "15 ml", "0,00307 $/ml", "0,046 $"],
        ["Tomates", "30 g", "0,0044 $/g", "0,132 $"],
        ["Boîte + serviette", "1+1", "—", "0,097 $"],
        ["TOTAL", "", "", "~1,13 $"],
    ])

    add_title(doc, "5. Ventes de la semaine — totaux")
    add_table(doc, ["POS ID", "Product", "Lun", "Mar", "Mer", "Jeu", "Ven", "TOTAL"], [
        ["101", "Burger classique", "25", "28", "22", "30", "35", "140"],
        ["102", "Burger bacon", "10", "12", "8", "15", "18", "63"],
        ["201", "Frites", "30", "35", "28", "40", "45", "178"],
        ["301", "Salade César", "8", "10", "12", "9", "14", "53"],
        ["401", "Bol chili", "15", "18", "20", "16", "22", "91"],
        ["501", "Cola", "40", "45", "38", "50", "55", "228"],
        ["502", "Eau", "12", "15", "10", "18", "20", "75"],
    ])
    add_p(doc, "Usage idéal approximatif bœuf burger (item Bœuf haché 80/20) : (140 + 63) boulettes × 77 g ≈ 15 631 g, plus chili et autres recettes.")

    add_title(doc, "6. Pertes — impact attendu")
    add_table(doc, ["Date", "Item", "Effet sur Usage Summary"], [
        ["Mardi", "Laitue 500 g", "Actual > Ideal sur laitue ; écart positif"],
        ["Mercredi", "Frites prep 400 g", "Actual prep frites augmente"],
        ["Jeudi", "Chili 800 g", "Perte prep ; shelf life expirée"],
        ["Vendredi", "Pain × 6", "Actual pain > ideal proportionnel"],
    ])

    add_title(doc, "7. Emplacements de stock — réponses")
    add_bullets_simple = [
        "Primary = emplacement de référence pour commandes et total consolidé.",
        "Secondary = où le stock est aussi physiquement (comptage par Location).",
        "Déplacer physiquement du stock entre armoires ne crée pas de transaction OC : seul le comptage par emplacement reflète la réalité.",
        "Tri Location à l'inventaire : feuille Réfrigérateur, puis Congélateur, etc.",
    ]
    for b in add_bullets_simple:
        doc.add_paragraph(b, style="List Bullet")

    add_title(doc, "8. Transferts entre établissements — flux")
    add_table(doc, ["Étape", "Magasin", "Menu", "Action"], [
        ["1", "Restaurant", "Stock Transfer → Request Stock", "Demande bœuf, pommes, cola à l'entrepôt"],
        ["2", "Entrepôt", "Approve Requests", "Approuve la demande"],
        ["3", "Entrepôt", "Fulfill Requisitions", "Expédie (partiel si stock insuffisant → backorder)"],
        ["Alt.", "Entrepôt", "Transfers → New", "Transfert direct sans réquisition"],
    ])
    add_p(doc, "Après fulfill : stock ↓ entrepôt, stock ↑ restaurant. Les deux magasins doivent partager les mêmes items (Item and Store Management).")

    add_title(doc, "9. Usage Summary — réponses types")
    add_bullets(doc_items := [
        "Plus grand écart $ : souvent bœuf haché ou pommes de terre (volume + ventes élevées).",
        "Ideal = 0, Actual > 0 : item sans recette liée ou Actualize activé (ex. huile si non liée aux frites).",
        "Food cost global : depend du mix ventes ; comparer catégorie Food seule.",
        "Drill-down Purchases : doit lister factures Distrib. Caraïbes, Boulangerie Pétion, etc.",
        "Drill-down Ideal : burger → products 101/102 ; frites → product 201.",
    ])
    for b in doc_items:
        doc.add_paragraph(b, style="List Bullet")

    add_title(doc, "10. Défis bonus — solutions")
    add_title(doc, "Défi A", 2)
    add_p(doc, "Inventory → Adjust Inventory → ouvrir l'inventaire d'ouverture ou le dernier → corriger pommes de terre 2 cs → 1 cs → Save.")
    add_title(doc, "Défi B", 2)
    add_p(doc, "Modifier recette Burger classique (+10 g bacon ou 1 tranche) → Sales → Rerun Sales Mix sur la période.")
    add_title(doc, "Défi C", 2)
    add_p(doc, "Miscellaneous → Item Management → Eau → décocher Active → Save. Historique ventes conservé.")
    add_title(doc, "Défi D", 2)
    add_p(doc, "Corriger prix bœuf sur item → Preferences → Inventory → Last Cost → ajuster comptage inventaire (1→0→1) → Summarize → Usage Summary.")

    add_title(doc, "Annexe — Sortie brute du calculateur")
    for line in calc_text.splitlines():
        doc.add_paragraph(line)

    doc.save(OUT)
    print(f"OK -> {OUT}")

if __name__ == "__main__":
    build()
