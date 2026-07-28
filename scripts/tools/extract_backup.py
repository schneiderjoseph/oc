from docx import Document

b = Document(r"E:\OC DOCS\Comprendre_Optimum_Control_backup.docx")

def extract_section(needle, max_paras=15):
    for i, p in enumerate(b.paragraphs):
        if needle in p.text:
            out = [f"=== {needle} @ {i} style={p.style.name} ==="]
            for j in range(i, min(i+max_paras, len(b.paragraphs))):
                p2 = b.paragraphs[j]
                t = p2.text.strip().replace("\n", " ")
                s = p2.style.name if p2.style else ""
                if not t: t = "(vide)"
                out.append(f"{j:4d}|{s:20s}|{t[:90]}")
            return out
    return [f"NOT FOUND: {needle}"]

needles = [
    "Vocabulaire essentiel",
    "Raccourcis clavier",
    "Volume et poids courants",
    "Méthodes de tri des feuilles",
    "Les quatre états",
    "Les achats de Joe",
    "Calcul du coût unitaire",
    "Deux types de base de données",
    "Détailler les pertes",
]
lines = []
for n in needles:
    lines.extend(extract_section(n))
    lines.append("")
open(r"E:\OC DOCS\backup_sections.txt", "w", encoding="utf-8").write("\n".join(lines))

# tables near quatre états
for ti, table in enumerate(b.tables):
    text = " | ".join(cell.text.strip() for row in table.rows for cell in row.cells)
    if any(w in text for w in ["Unlinked", "Linked", "Valid", "Mismatch", "Ignored"]):
        open(r"E:\OC DOCS\tables_quatre.txt", "w", encoding="utf-8").write(f"Table {ti}:\n{text}\n")
