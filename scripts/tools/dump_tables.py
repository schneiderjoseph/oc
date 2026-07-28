from docx import Document

b = Document(r"E:\OC DOCS\Comprendre_Optimum_Control_backup.docx")
lines = []
for ti, table in enumerate(b.tables):
    rows = []
    for row in table.rows:
        rows.append(" | ".join(c.text.strip().replace("\n"," ") for c in row.cells))
    text = "\n".join(rows)
    if len(text.strip()) < 5:
        continue
    # context: paragraph before table is hard; just dump all tables with index
    lines.append(f"=== TABLE {ti} ({len(table.rows)} rows) ===")
    lines.append(text)
    lines.append("")
open(r"E:\OC DOCS\all_tables.txt", "w", encoding="utf-8").write("\n".join(lines))
print(len(b.tables), "tables")
