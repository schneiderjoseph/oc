from docx import Document
d = Document(r"E:\OC DOCS\Comprendre_Optimum_Control_V2.docx")
lines = []
for needle in ["Détailler les pertes", "Taxes, ajustements", "Sales Export", "Listes chaudes", "Niveaux Par des items", "Créer et déployer"]:
    for i, p in enumerate(d.paragraphs):
        if needle in p.text:
            lines.append(f"\n=== {needle} @ {i} ===")
            for j in range(i, min(i+6, len(d.paragraphs))):
                t = d.paragraphs[j].text.strip()[:70] if d.paragraphs[j].text.strip() else "(vide)"
                sn = d.paragraphs[j].style.name if d.paragraphs[j].style else "X"
                lines.append(f"  {j}|{sn}|{t}")
            break
open(r"E:\OC DOCS\ranges.txt", "w", encoding="utf-8").write("\n".join(lines))
