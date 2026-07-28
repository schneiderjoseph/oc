#!/usr/bin/env python3
"""Detect headings with no body content below them."""
from docx import Document

doc = Document(r"E:\OC DOCS\Comprendre_Optimum_Control_V2.docx")

def heading_level(style_name):
    if not style_name or not style_name.startswith("Heading"):
        return 99
    n = style_name.replace("Heading ", "")
    return int(n) if n.isdigit() else 99

problems = []
for i, p in enumerate(doc.paragraphs):
    sn = p.style.name if p.style else ""
    t = p.text.strip()
    if not t or not sn.startswith("Heading"):
        continue
    level = heading_level(sn)
    has_body = False
    next_info = None
    for j in range(i + 1, min(i + 12, len(doc.paragraphs))):
        q = doc.paragraphs[j]
        qt = q.text.strip()
        qs = q.style.name if q.style else ""
        if not qt:
            continue
        if qs.startswith("Heading"):
            ql = heading_level(qs)
            if ql <= level:
                next_info = f"[{qs}] {qt[:55]}"
                break
            # subheading counts as content
            has_body = True
            break
        has_body = True
        next_info = f"OK: {qt[:55]}"
        break
    if not has_body:
        problems.append((i, sn, t, next_info))

lines = [f"TITRES SANS TEXTE EN DESSOUS: {len(problems)}\n"]
for idx, sn, t, nxt in problems:
    lines.append(f"[{idx}] {sn}")
    lines.append(f"  {t}")
    lines.append(f"  -> suivant: {nxt}\n")

open(r"E:\OC DOCS\empty_headings.txt", "w", encoding="utf-8").write("\n".join(lines))
print(f"written {len(problems)} problems to empty_headings.txt")
