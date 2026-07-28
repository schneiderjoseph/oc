from docx import Document
b = Document(r"E:\OC DOCS\Comprendre_Optimum_Control_backup.docx")
lines = []
for needle in ["Vocabulaire", "Raccourcis", "Les quatre états", "Méthodes de tri"]:
    for i, p in enumerate(b.paragraphs):
        if needle in p.text:
            lines.append(f"BACKUP {needle} @ {i}")
            for j in range(i, min(i+6, len(b.paragraphs))):
                t = b.paragraphs[j].text.strip()[:65] if b.paragraphs[j].text.strip() else "(vide)"
                lines.append(f"  {j}|{t}")
            break
open(r"E:\OC DOCS\bkp.txt", "w", encoding="utf-8").write("\n".join(lines))
