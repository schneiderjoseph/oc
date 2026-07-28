from docx import Document

def map_tables(path, label):
    doc = Document(path)
    lines = [f"=== {label}: {len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs ==="]
    # python-docx doesn't give table position in body easily; use element order
    from docx.oxml.ns import qn
    body = doc.element.body
    ti = 0
    pi = 0
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            if pi < len(doc.paragraphs):
                t = doc.paragraphs[pi].text.strip()[:60]
                s = doc.paragraphs[pi].style.name if doc.paragraphs[pi].style else ""
                if s.startswith("Heading") or any(k in t for k in ["Vocabulaire", "quatre états", "Volume et poids", "Méthodes de tri", "achats de Joe", "coût unitaire", "Deux types", "Raccourcis"]):
                    lines.append(f"  P{pi}|{s}|{t}")
            pi += 1
        elif tag == "tbl":
            table = doc.tables[ti]
            preview = " | ".join(c.text.strip()[:30] for c in table.rows[0].cells)
            lines.append(f"  >>> TABLE {ti}: {preview}")
            ti += 1
    return lines

open(r"E:\OC DOCS\table_positions.txt", "w", encoding="utf-8").write(
    "\n".join(map_tables(r"E:\OC DOCS\Comprendre_Optimum_Control_backup.docx", "BACKUP") + ["\n"] + map_tables(r"E:\OC DOCS\Comprendre_Optimum_Control_V2.docx", "V2"))
)
