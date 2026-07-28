from docx import Document
b=Document(r'E:\OC DOCS\Comprendre_Optimum_Control_backup.docx')
lines=[]
for i,p in enumerate(b.paragraphs):
    if 'Transfert de stock' in p.text or 'Emplacements de stockage' in p.text:
        for j in range(i, min(i+6,len(b.paragraphs))):
            lines.append(b.paragraphs[j].text)
open(r'E:\OC DOCS\storage_full.txt','w',encoding='utf-8').write('\n---\n'.join(lines))
