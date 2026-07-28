#!/usr/bin/env python3
"""Rapport Excel audit La Réserve + contrat retouché (souligné)."""
from __future__ import annotations

import csv
import re
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

BASE = Path(r"E:\OC DOCS")
EXPORT = Path(r"E:\EXPORT")
CSV_PATH = EXPORT / "bar.csv"
OUT_XLSX = BASE / "LaReserve_Audit_OC_Critique.xlsx"
OUT_DOCX = BASE / "Contrat_Prestation_OC_LaReserve_Retouche_Josep.docx"

# Palette visuelle — impact direction
C_RED_DARK = "C00000"
C_RED_LIGHT = "FFC7CE"
C_RED_MID = "FF6B6B"
C_ORANGE = "ED7D31"
C_ORANGE_LIGHT = "FCE4D6"
C_YELLOW_LIGHT = "FFF2CC"
C_GREEN_LIGHT = "E2EFDA"
C_HEADER_DARK = "833C0C"
C_HEADER_P1 = "C00000"
C_HEADER_P2 = "ED7D31"
C_HEADER_P3 = "BF8F00"
C_HEADER_NEUTRAL = "1F4E79"
C_ALERT_BG = "F4CCCC"
C_TITLE_BG = "5B0F0F"

THIN = Side(style="thin", color="999999")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

CONSULTATION_NOTE = (
    "Lors de la consultation sur site du système Optimum Control avec "
    "**M. Ricardo** et **Mme Jeannine**, les données inventaire (items, emplacements, "
    "unités de mesure, conversions et valorisations) ont été consultées et extraites "
    "à des fins d'analyse uniquement. Ce rapport mesure la **gravité** des anomalies "
    "constatées et propose un plan de correction sur 10 mois — sans engagement contractuel "
    "à ce stade."
)
SRC_DOCX = BASE / "Contrat Prestation Services_Optimum Control_Juillet 2026.docx"


def parse_val(s: str) -> float:
    s = (s or "").replace("$", "").replace(",", "").strip()
    try:
        return float(s)
    except ValueError:
        return 0.0


def load_rows():
    with CSV_PATH.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def classify_rows(rows):
    """Return list of dicts with issue metadata."""
    issues = []
    by_item = defaultdict(list)
    for r in rows:
        by_item[r["textBox11"].strip().upper()].append(r)

    seen_dup_report = set()
    for r in rows:
        name = r["textBox11"].strip()
        loc = r["textBox29"].strip()
        u1, u2, u3 = r["textBox20"].strip(), r["textBox21"].strip(), r["textBox22"].strip()
        conv = r["textBox31"].strip()
        val = parse_val(r["textBox13"])
        key = (name, loc)

        # P1 — bloquant food cost / inventaire
        if re.search(r"^\.{3,}|^\?|^\.\.\.|^M$|^E\d+$", name) or "????" in name:
            issues.append(_issue("P1", "Nom invalide / placeholder", name, loc, u1, u2, u3, conv, val,
                                 "Item à renommer ou supprimer — fausse les rapports et exports."))
        if name.upper() == "BEEF" and "bottle" in (u1.lower(), u3.lower()):
            issues.append(_issue("P1", "Produit vs unité incohérent", name, loc, u1, u2, u3, conv, val,
                                 "Bœuf en bottle/ml — coût et usage idéal faux."))
        if "GANT" in name.upper() and u3.lower() == "gal":
            issues.append(_issue("P1", "Conversion absurde", name, loc, u1, u2, u3, conv, val,
                                 "Gants en gallons — valeur stock aberrante."))
        if not conv and u1.lower() in ("batch", "marmite"):
            issues.append(_issue("P1", "Prep/batch sans conversion", name, loc, u1, u2, u3, conv, val,
                                 "Recette batch inutilisable pour usage idéal."))
        if "mg / bottle" in conv.lower():
            issues.append(_issue("P1", "mg au lieu de ml", name, loc, u1, u2, u3, conv, val,
                                 "Erreur d'unité — coût portion ×1000."))
        if conv.strip().upper() in ("MUSCADOR", "LIQUOR") or (conv == "1.12" and "DIMPLE" in name.upper()):
            issues.append(_issue("P1", "Conversion non numérique", name, loc, u1, u2, u3, conv, val,
                                 "Champ conversion = texte ou valeur incohérente."))
        if val >= 15_000:
            issues.append(_issue("P1", "Valeur inventaire aberrante", name, loc, u1, u2, u3, conv, val,
                                 f"Stock valorisé {val:,.0f} $ — vérifier case size / conversion."))

        # P2 — bar / alcools (stratégique marge)
        if loc == "Bar":
            spirit_kw = ("GIN", "WHISK", "COGNAC", "HENNESSY", "RUM", "VODKA", "TEQUILA", "APEROL", "PINCH", "DSP")
            if any(k in name.upper() for k in spirit_kw):
                ok = u1.lower() == "bottle" and u3.lower() in ("oz (fl)", "ml")
                if not ok:
                    issues.append(_issue("P2", "Spiritueux bar — mauvaise UOM", name, loc, u1, u2, u3, conv, val,
                                         "Configurer bottle → oz (fl) ou ml (750) pour coût cocktail."))

        if loc == "Cooler" and u1.lower() == "ea":
            if any(k in name.upper() for k in ("CHATEAU", "VIN", "WINE", "MERLOT", "MOET", "PROSECCO")):
                issues.append(_issue("P2", "Vin en 'ea' au cooler", name, loc, u1, u2, u3, conv, val,
                                     "Vins doivent être en bottle + ml pour coût verre."))

        if "*" in conv and "/" not in conv:
            issues.append(_issue("P2", "Conversion format 1*XX", name, loc, u1, u2, u3, conv, val,
                                 "Format non standard OC — risque erreur à l'import / comptage."))

        if conv.strip() == "1*1" and "ml" in u3.lower():
            issues.append(_issue("P2", "Spiritueux 1*1 suspect", name, loc, u1, u2, u3, conv, val,
                                 "Taille bouteille non définie (attendu 750 ml, etc.)."))

        # P3 — doublons multi-emplacements
        item_key = name.upper()
        if item_key not in seen_dup_report and len(by_item[item_key]) > 1:
            locs = {x["textBox29"].strip() for x in by_item[item_key]}
            if len(locs) > 1:
                configs = {
                    (x["textBox20"], x["textBox21"], x["textBox22"], x["textBox31"]) for x in by_item[item_key]
                }
                prio = "P2" if len(locs) >= 4 or len(configs) > 1 else "P3"
                issues.append(_issue(
                    prio, f"Doublon — {len(locs)} emplacements",
                    name, ", ".join(sorted(locs)), u1, u2, u3,
                    f"{len(configs)} config(s)", sum(parse_val(x["textBox13"]) for x in by_item[item_key]),
                    "Amalgamate + Primary Location — sinon stock et coût double comptés.",
                ))
                seen_dup_report.add(item_key)

    # dédup exact rows
    uniq = {}
    for i in issues:
        k = (i["Priorité"], i["Type"], i["Item"], i["Emplacement"])
        if k not in uniq or i["Valeur $"] > uniq[k]["Valeur $"]:
            uniq[k] = i
    return sorted(uniq.values(), key=lambda x: (x["Priorité"], -x["Valeur $"], x["Item"]))


def _issue(prio, typ, name, loc, u1, u2, u3, conv, val, impact):
    return {
        "Priorité": prio,
        "Type": typ,
        "Item": name,
        "Emplacement": loc,
        "UOM achat": u1,
        "UOM split": u2,
        "UOM recette": u3,
        "Conversion": conv,
        "Valeur $": round(val, 2) if isinstance(val, (int, float)) else val,
        "Impact métier": impact,
    }


def style_header(ws, row=1, fg=C_HEADER_NEUTRAL):
    fill = PatternFill("solid", fgColor=fg)
    font = Font(bold=True, color="FFFFFF", size=11)
    for cell in ws[row]:
        cell.fill = fill
        cell.font = font
        cell.border = BORDER
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def fill_row(ws, row_idx, fg, bold=False, font_color="000000"):
    for cell in ws[row_idx]:
        cell.fill = PatternFill("solid", fgColor=fg)
        cell.border = BORDER
        cell.alignment = Alignment(vertical="top", wrap_text=True)
        if bold:
            cell.font = Font(bold=True, color=font_color)


def style_priority_rows(ws, headers, data, prio_col="Priorité"):
    prio_idx = headers.index(prio_col) + 1 if prio_col in headers else None
    for r_idx, row in enumerate(data, start=2):
        prio = row.get(prio_col, "")
        if prio == "P1":
            fill_row(ws, r_idx, C_RED_LIGHT, font_color=C_RED_DARK)
            if prio_idx:
                ws.cell(r_idx, prio_idx).font = Font(bold=True, color="FFFFFF")
                ws.cell(r_idx, prio_idx).fill = PatternFill("solid", fgColor=C_RED_DARK)
        elif prio == "P2":
            fill_row(ws, r_idx, C_ORANGE_LIGHT)
            if prio_idx:
                ws.cell(r_idx, prio_idx).font = Font(bold=True, color="FFFFFF")
                ws.cell(r_idx, prio_idx).fill = PatternFill("solid", fgColor=C_ORANGE)
        elif prio == "P3":
            fill_row(ws, r_idx, C_YELLOW_LIGHT)
            if prio_idx:
                ws.cell(r_idx, prio_idx).font = Font(bold=True)
                ws.cell(r_idx, prio_idx).fill = PatternFill("solid", fgColor=C_HEADER_P3)


def write_issue_sheet(wb, title, header_color, data, issues_fallback):
    ws = wb.create_sheet(title[:31])
    headers = list(data[0].keys()) if data else list(issues_fallback[0].keys())
    ws.append(headers)
    style_header(ws, fg=header_color)
    for row in data:
        ws.append([row[h] for h in headers])
    style_priority_rows(ws, headers, data)
    ws.freeze_panes = "A2"
    autosize(ws)
    return ws


def autosize(ws, max_width=55):
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        width = max(12, min(max_width, max(len(str(c.value or "")) for c in col) + 2))
        ws.column_dimensions[letter].width = width


def write_excel(rows, issues):
    wb = Workbook()
    ws = wb.active
    ws.title = "Résumé exécutif"

    items = {r["textBox11"].strip() for r in rows}
    locs = Counter(r["textBox29"].strip() for r in rows)
    p1 = sum(1 for i in issues if i["Priorité"] == "P1")
    p2 = sum(1 for i in issues if i["Priorité"] == "P2")
    p3 = sum(1 for i in issues if i["Priorité"] == "P3")
    total_val = sum(parse_val(r["textBox13"]) for r in rows)
    star = sum(1 for r in rows if "*" in r["textBox31"] and "/" not in r["textBox31"])

    # --- Bandeau titre ---
    ws.merge_cells("A1:D1")
    t = ws["A1"]
    t.value = "AUDIT OPTIMUM CONTROL — LA RÉSERVE"
    t.font = Font(bold=True, size=18, color="FFFFFF")
    t.fill = PatternFill("solid", fgColor=C_TITLE_BG)
    t.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    ws.merge_cells("A2:D2")
    sub = ws["A2"]
    sub.value = "Rapport préliminaire de gravité des données — à l'attention de la direction"
    sub.font = Font(bold=True, size=12, color=C_RED_DARK)
    sub.fill = PatternFill("solid", fgColor=C_ALERT_BG)
    sub.alignment = Alignment(horizontal="center", wrap_text=True)
    ws.row_dimensions[2].height = 28

    # --- Contexte consultation (sans chemin fichier) ---
    ws.merge_cells("A4:D8")
    ctx = ws["A4"]
    ctx.value = (
        "Contexte : lors de la consultation sur site du système Optimum Control avec "
        "M. Ricardo et Mme Jeannine, les données inventaire (items, emplacements, unités, "
        "conversions et valorisations) ont été consultées et extraites afin d'en analyser "
        "la cohérence et la fiabilité pour la gestion (food cost, bar, inventaires, rapports).\n\n"
        "Ce document présente la gravité des anomalies détectées et un plan de correction "
        "sur 10 mois. Il sert de base factuelle avant toute proposition d'accompagnement."
    )
    ctx.font = Font(size=11)
    ctx.fill = PatternFill("solid", fgColor="F2F2F2")
    ctx.alignment = Alignment(wrap_text=True, vertical="top")
    ctx.border = BORDER
    ws.row_dimensions[4].height = 95

    ws["A9"] = f"Date du rapport"
    ws["B9"] = date.today().strftime("%d/%m/%Y")
    ws["A9"].font = Font(bold=True)

    # --- ALERTE rouge synthèse ---
    row_alert = 11
    ws.merge_cells(f"A{row_alert}:D{row_alert}")
    alert = ws[f"A{row_alert}"]
    alert.value = (
        f"ALERTE — {p1} anomalies CRITIQUES (P1) · {p2} importantes (P2) · "
        f"{p3} à planifier (P3)"
    )
    alert.font = Font(bold=True, size=13, color="FFFFFF")
    alert.fill = PatternFill("solid", fgColor=C_RED_DARK)
    alert.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[row_alert].height = 30

    # --- Tableau indicateurs ---
    start = 13
    metrics = [
        ("Indicateur", "Valeur", "Niveau", ""),
        ("Anomalies P1 — CRITIQUE (fiabilité compromise)", p1, "URGENT", "Corriger en priorité absolue"),
        ("Anomalies P2 — IMPORTANT (bar, vins, doublons)", p2, "ÉLEVÉ", "Impact marge & stock"),
        ("Anomalies P3 — À planifier", p3, "MOYEN", "Consolidation progressive"),
        ("Fiches item × emplacement analysées", len(rows), "INFO", ""),
        ("Items distincts", len(items), "INFO", ""),
        ("Emplacements de stockage", len(locs), "INFO", ""),
        ("Conversions format 1*XX (non standard)", star, "ÉLEVÉ", "Risque erreur comptage"),
        ("Valeur stock affichée (total export)", f"{total_val:,.0f} $US", "À AUDITER", "Plusieurs lignes aberrantes"),
    ]
    for i, row_data in enumerate(metrics):
        r = start + i
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(r, c, val)
            cell.border = BORDER
            if i == 0:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor=C_HEADER_DARK)
            elif c == 3 and val == "URGENT":
                cell.fill = PatternFill("solid", fgColor=C_RED_DARK)
                cell.font = Font(bold=True, color="FFFFFF")
            elif c == 3 and val == "ÉLEVÉ":
                cell.fill = PatternFill("solid", fgColor=C_ORANGE)
                cell.font = Font(bold=True, color="FFFFFF")
            elif c == 3 and val == "MOYEN":
                cell.fill = PatternFill("solid", fgColor=C_HEADER_P3)
                cell.font = Font(bold=True)
            elif c == 1 and i > 0 and "P1" in str(row_data[0]):
                cell.fill = PatternFill("solid", fgColor=C_RED_LIGHT)
                cell.font = Font(bold=True, color=C_RED_DARK)
            elif c == 2 and i == 1:
                cell.fill = PatternFill("solid", fgColor=C_RED_LIGHT)
                cell.font = Font(bold=True, size=12, color=C_RED_DARK)

    msg_row = start + len(metrics) + 2
    ws.merge_cells(f"A{msg_row}:D{msg_row + 2}")
    msg = ws[f"A{msg_row}"]
    msg.value = (
        "MESSAGE DIRECTION\n\n"
        "Tant que les anomalies P1 (et une partie des P2) ne sont pas corrigées, "
        "les rapports Usage Summary, les % food cost et les coûts bar/cocktails "
        "NE PEUVENT PAS être utilisés pour piloter les prix, les achats ou le contrôle des pertes.\n\n"
        "Recommandation : plan structuré sur 10 mois (voir onglet « Plan 10 mois »)."
    )
    msg.font = Font(bold=True, size=11, color=C_RED_DARK)
    msg.fill = PatternFill("solid", fgColor=C_ALERT_BG)
    msg.alignment = Alignment(wrap_text=True, vertical="top")
    msg.border = BORDER
    ws.row_dimensions[msg_row].height = 75

    ws.column_dimensions["A"].width = 44
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 38

    # --- Feuilles priorités ---
    write_issue_sheet(wb, "P1 — CRITIQUE", C_HEADER_P1,
                      [i for i in issues if i["Priorité"] == "P1"], issues)
    write_issue_sheet(wb, "P2 — Important", C_HEADER_P2,
                      [i for i in issues if i["Priorité"] == "P2"], issues)
    write_issue_sheet(wb, "P3 — Planifier", C_HEADER_P3,
                      [i for i in issues if i["Priorité"] == "P3"], issues)

    # --- Top valeurs ---
    ws = wb.create_sheet("Top valeurs stock")
    ws.append(["Rang", "Item", "Emplacement", "UOM", "Conversion", "Valeur $US", "Alerte"])
    style_header(ws, fg=C_HEADER_P1)
    for rank, r in enumerate(sorted(rows, key=lambda x: parse_val(x["textBox13"]), reverse=True)[:40], 1):
        v = parse_val(r["textBox13"])
        if v >= 25_000:
            alert_txt = "CRITIQUE — vérifier immédiatement"
            fg = C_RED_LIGHT
        elif v >= 10_000:
            alert_txt = "ÉLEVÉ — probable erreur conversion"
            fg = C_ORANGE_LIGHT
        elif v >= 5_000:
            alert_txt = "À vérifier"
            fg = C_YELLOW_LIGHT
        else:
            alert_txt = ""
            fg = None
        row_idx = ws.max_row + 1
        ws.append([
            rank, r["textBox11"].strip(), r["textBox29"].strip(),
            f'{r["textBox20"]}/{r["textBox21"]}/{r["textBox22"]}',
            r["textBox31"].strip(), v, alert_txt,
        ])
        if fg:
            fill_row(ws, row_idx, fg, bold=(v >= 15_000), font_color=C_RED_DARK if v >= 25_000 else "000000")
    ws.freeze_panes = "A2"
    autosize(ws)

    # --- Emplacements ---
    ws = wb.create_sheet("Par emplacement")
    ws.append(["Emplacement", "Nb fiches", "% du total", "Charge data"])
    style_header(ws)
    for loc, cnt in locs.most_common():
        pct = round(100 * cnt / len(rows), 1)
        charge = "Élevée" if cnt > 400 else "Moyenne" if cnt > 100 else "Normale"
        row_idx = ws.max_row + 1
        ws.append([loc, cnt, pct, charge])
        if charge == "Élevée":
            fill_row(ws, row_idx, C_ORANGE_LIGHT)
    autosize(ws)

    # --- Plan 10 mois ---
    ws = wb.create_sheet("Plan 10 mois")
    plan = [
        ("Phase", "Mois", "Priorité", "Objectif", "Livrables"),
        ("1 — Audit & stabilisation", "1–2", "P1", "Sauvegardes + rapport validé", "Rapport écrit, liste fusions validées"),
        ("2 — Bar & alcools", "3–4", "P2", "Spiritueux/vins, conversions bouteille", "Coût cocktail/vin fiable"),
        ("3 — Doublons & amalgamates", "5–7", "P2/P3", "Doublons multi-emplacements", "Stock consolidé par item"),
        ("4 — POS & ventes", "7–8", "P2", "Pending Sales, Sales Mix, Daily Sales", "0 Unlinked, Z caisse"),
        ("5 — Formation & usage", "9–10", "—", "Inventaire, Usage Summary, procédures", "Guide + 2 sessions équipe"),
    ]
    for i, row in enumerate(plan):
        ws.append(row)
        if i == 0:
            style_header(ws, fg=C_HEADER_NEUTRAL)
        else:
            prio = row[2]
            if prio == "P1":
                fill_row(ws, i + 1, C_RED_LIGHT)
            elif prio.startswith("P2"):
                fill_row(ws, i + 1, C_ORANGE_LIGHT)
    autosize(ws)

    wb.save(OUT_XLSX)
    print(f"OK -> {OUT_XLSX}")


def add_mixed_paragraph(doc, parts):
    """parts: list of (text, underline_bool)"""
    p = doc.add_paragraph()
    for text, ul in parts:
        run = p.add_run(text)
        if ul:
            run.underline = WD_UNDERLINE.SINGLE
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xAA)


def build_contract():
    doc = Document()
    title = doc.add_heading("CONTRAT D'ASSISTANCE FONCTIONNELLE OPTIMUM CONTROL", 0)
    doc.add_paragraph(
        "Version retouchée — propositions du Prestataire (Josep). "
        "Les passages soulignés sont des ajouts ou modifications par rapport au projet La Réserve."
    ).runs[0].italic = True

    sections = [
        ("1. Les Parties", False, [
            ("Entre les soussignés :\n", False),
            ("Le Client : La Réserve / [raison sociale complète], situé à [adresse Port-au-Prince].\n", True),
            ("Le Prestataire : [Nom du technicien / Spécialiste IT], [Numéro Siret/Fiscal / NIF], situé à [adresse].\n", False),
        ]),
        ("2. Objet de la Mission", False, [
            ("Le présent contrat définit les conditions d'une intervention ", False),
            ("sur une durée de dix (10) mois", True),
            (" visant l'audit, la correction progressive, le nettoyage des données, la configuration des flux de caisse, la formation de l'équipe et la fourniture de documentation sur le système Optimum Control du Client.\n", False),
        ]),
        ("3. Description des Prestations (Périmètre par phases)", True, [
            ("Le Prestataire réalise les tâches suivantes, ", False),
            ("selon le plan en annexe technique (rapport Excel et calendrier 10 mois)", True),
            (" :\n", False),
            ("• Phase 1 (mois 1–2) — Audit, sauvegarde, rapport écrit des anomalies P1/P2, plan de correction validé par le Client.\n", True),
            ("• Phase 2 (mois 3–4) — Corrections critiques bar/alcools, unités de mesure et conversions bouteille.\n", True),
            ("• Phase 3 (mois 5–7) — Amalgamation des doublons ", False),
            ("(liste préalable validée par écrit par le Client avant toute fusion)", True),
            (", réorganisation des emplacements.\n", False),
            ("• Phase 4 (mois 7–8) — Ventes POS : Pending Sales, Sales Mix, Daily Sales.\n", True),
            ("• Phase 5 (mois 9–10) — Inventaire, Usage Summary, formation et guide utilisateur.\n", True),
            ("• Nettoyage fonctionnel : doublons, erreurs d'unités, fiches fournisseurs ou recettes erronées (via interface OC uniquement).\n", False),
            ("• Rapport d'audit écrit et guide utilisateur personnalisé — livrables obligatoires.\n", False),
            ("• Formation : ", False),
            ("deux (2) sessions pratiques", True),
            (" minimum (saisie + inventaire/rapports).\n", False),
        ]),
        ("4. Modalités de Réalisation (Sur site et à distance)", False, [
            ("Mode mixte sur ", False),
            ("quatre (4) jours par semaine, six à huit (6–8) heures par jour", True),
            (" :\n", False),
            ("• ", False),
            ("Deux (2) jours sur site minimum", True),
            (" (audit, inventaire, formation, validation).\n", False),
            ("• ", False),
            ("Deux (2) jours à distance au choix du Prestataire", True),
            (" si le travail (données, rapport, paramétrage) ne nécessite pas de présence physique.\n", False),
            ("• Connexion sécurisée : TeamViewer / AnyDesk — accès fourni par le Client aux horaires convenus.\n", True),
        ]),
        ("5. Limites Fermes de la Mission (Exclusions)", False, [
            ("Aucun accès ni modification du code source Optimum Control.\n", False),
            ("Aucune modification du schéma de base de données (tables, index, SQL d'altération structurelle).\n", False),
            ("Corrections exclusivement via l'interface OC et outils d'import standard.\n", False),
            ("Hors périmètre sauf avenant : ", False),
            ("re-saisie complète de toutes les recettes, développement logiciel, inventaire physique à la place du personnel Client.\n", True),
            ("Volume indicatif audit initial : ", False),
            ("~1 200 items, 9 emplacements, anomalies P1 documentées dans le rapport Excel.", True),
            ("\n", False),
        ]),
        ("6. Délai d'Exécution et Durée du Contrat", False, [
            ("Date de début : [Date].\n", False),
            ("Durée : ", False),
            ("dix (10) mois", True),
            (" à compter de la date de début.\n", False),
            ("Charge incluse : ", False),
            ("jusqu'à cent dix (110) heures par mois", True),
            (" (4 j × ~7 h × 4 semaines). Au-delà : facturation selon article 7 ou avenant.\n", True),
            ("Fin de mission : validation des livrables de la phase 5 et attestation de formation signée.\n", True),
        ]),
        ("7. Tarifs et Modalités de Paiement", True, [
            ("Le Client formule une ", False),
            ("proposition de forfait mensuel", True),
            (" couvrant l'ensemble des prestations (présence, audit, corrections, formation, déplacements locaux).\n", False),
            ("Montant mensuel proposé par le Client : ", False),
            ("[Montant $US ou HTG / mois]", True),
            (" — à compléter par La Réserve.\n", False),
            ("Paiement : ", False),
            ("mensuel, à réception de facture, au plus tard quinze (15) jours", True),
            (" après chaque mois civil.\n", False),
            ("Premier paiement : ", False),
            ("après remise du rapport d'audit phase 1 (fin du mois 2) ou à la signature, selon accord.", True),
            ("\n", False),
            ("Heures supplémentaires non incluses : ", False),
            ("[taux horaire $US] / heure", True),
            (" avec accord écrit préalable.\n", False),
        ]),
        ("8. Responsabilités et Sauvegarde", False, [
            ("Le Client réalise une ", False),
            ("sauvegarde complète (File → Backup Data) avant toute intervention", True),
            (" et confirme par écrit (e-mail) la date de sauvegarde.\n", False),
            ("Seconde sauvegarde obligatoire ", False),
            ("avant toute vague d'amalgamation d'items (phase 3).", True),
            ("\n", False),
            ("Responsabilité du Prestataire limitée aux traitements validés et aux fonctionnalités natives du logiciel.\n", False),
            ("Le Client valide par écrit la liste des fusions d'items avant exécution.\n", True),
        ]),
        ("9. Confidentialité", False, [
            ("Strictement confidentiel : données commerciales, prix, recettes, marges.\n", False),
        ]),
        ("10. Option de Renouvellement", False, [
            ("Renouvellement possible par avenant écrit (suivi mensuel, maintenance données) — ", False),
            ("non tacite", True),
            (".\n", False),
        ]),
        ("11. Litiges", False, [
            ("Droit en vigueur en République d'Haïti. Tribunal compétent : ", False),
            ("Port-au-Prince", True),
            (".\n", False),
        ]),
        ("Annexe A — Critères d'acceptation (fin de mission)", True, [
            ("• Rapport d'audit et plan 10 mois remis et signés.\n", True),
            ("• Anomalies P1 du rapport initial corrigées ou documentées avec accord Client.\n", True),
            ("• Spiritueux bar : configuration bottle + oz/ml sur les items prioritaires.\n", True),
            ("• Amalgamations réalisées selon liste validée.\n", True),
            ("• Ventes : Pending Sales traités (objectif 0 Unlinked sur période test).\n", True),
            ("• Deux sessions de formation réalisées + guide utilisateur remis.\n", True),
            ("• Usage Summary généré sur une période test sans blocage majeur.\n", True),
        ]),
    ]

    for heading, heading_ul, parts in sections:
        h = doc.add_heading(heading, level=1)
        if heading_ul:
            for run in h.runs:
                run.underline = WD_UNDERLINE.SINGLE
        add_mixed_paragraph(doc, parts)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Fait en deux exemplaires à Port-au-Prince, le [Date].\n\n")
    p.add_run("Signature Client — Lu et approuvé\n\n")
    p.add_run("Signature Prestataire — Lu et approuvé\n")

    doc.save(OUT_DOCX)
    print(f"OK -> {OUT_DOCX}")


def main():
    rows = load_rows()
    issues = classify_rows(rows)
    write_excel(rows, issues)
    p1 = sum(1 for i in issues if i["Priorité"] == "P1")
    p2 = sum(1 for i in issues if i["Priorité"] == "P2")
    print(f"Anomalies: P1={p1} P2={p2} total={len(issues)}")


if __name__ == "__main__":
    main()
