#!/usr/bin/env python3
"""Parcours linéaire : tutoriels OC + configuration complète (Store, Setup, Preferences)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from exercice_locale import (
    ADRESSE, DEVISE, DEVISE_COURTE, NOTE_DEVISE, SEMAINE_RAPPORT, TAXES, TAXES_DOC, TAXES_NOTE, TAX_HEADERS, tax_table_rows,
    TAX_GROUPS_NOTE, TAX_GROUP_HEADERS, tax_group_rows, TAX_GROUP_CODE,
    FOURNISSEURS, NOM_DISTRIB, NOM_BOULANGERIE, NOM_EMBALLAGES, NOM_BOISSONS, COMPTES_GL,
    ACCOUNT_TYPES_DOC, CATEGORIES, INVENTORY_GROUPS, SALES_GROUPS, SETUP_CATEGORIES_NOTE,
    SUPPLIERS_NOTE, SUPPLIER_HEADERS, supplier_table_rows,
    INVENTORY_PREFS_NOTE, INVENTORY_PREFS_HEADERS, inventory_prefs_rows,
    POS_PREFS_NOTE, POS_PREFS_HEADERS, pos_prefs_rows,
    POS_SPEC_HEADERS, pos_spec_rows,
    TILL_TAPE_NOTE, TILL_TAPE_TWO_STEP_NOTE, TILL_TAPE_POS_NOTE,
    TILL_TAPE_LIST_FLOW, TILL_TAPE_SALE_FLOW, TILL_TAPE_VALIDATE,
    TILL_TAPE_QTY_HEADERS, till_tape_mardi_rows, DAILY_SALES_GROSS_NOTE,
    PENDING_SALES_NOTE, PENDING_SALES_UNLINK_NOTE, PENDING_SALES_FLOW, PENDING_SALES_COMPTOIR_NOTE,
    PENDING_SALES_VALIDATE, PENDING_SALES_STATES_HEADERS, pending_sales_state_rows,
    WASTE_NOTE, WASTE_FLOW, WASTE_VALIDATE, WASTE_TABLE_HEADERS, waste_rows,
    DAILY_SALES_NOTE, DAILY_SALES_SCREEN_NOTE, DAILY_SALES_GROSS_NOTE, DAILY_SALES_DISTRIB_NOTE,
    DAILY_SALES_FLOW, DAILY_SALES_VALIDATE,
    DAILY_SALES_FIELDS_HEADERS, daily_sales_field_rows,
    DAILY_SALES_OPS_HEADERS, daily_sales_ops_rows,
    DAILY_SALES_TABLE_HEADERS, daily_sales_rows,
    CLOSING_INVENTORY_NOTE, CLOSING_INVENTORY_FLOW, CLOSING_INVENTORY_VALIDATE,
    CLOSING_INVENTORY_QTY_HEADERS, closing_inventory_qty_rows,
    USAGE_SUMMARY_NOTE, USAGE_SUMMARY_FORMULA, USAGE_SUMMARY_FLOW, USAGE_SUMMARY_QUESTIONS, USAGE_SUMMARY_VALIDATE,
    BACKUP_NOTE, BACKUP_FLOW,
    CSV_IMPORT_NOTE, CSV_IMPORT_HEADERS, csv_import_map_rows,
    ITEM_CREATE_NOTE, ITEM_CREATE_HEADERS, item_create_rows,
    ITEMS_FIRST3_NOTE, ITEMS_FIRST3_CORE_HEADERS, items_first3_core_rows,
    ITEMS_FIRST3_CASE_HEADERS, items_first3_case_rows,
    ITEM_MULTI_CASESIZE_NOTE,
    ITEM_MULTI_CASESIZE_EXAMPLE_NOTE, ITEM_MULTI_CASESIZE_EXAMPLE_HEADERS,
    items_multicasize_example_rows,
    ITEMS_REST13_NOTE, ITEMS_DETAIL_CORE_HEADERS, items_rest13_core_rows,
    ITEMS_DETAIL_CASE_HEADERS, items_rest13_case_rows,
    ITEMS_LAST5_NOTE, items_last5_core_rows, items_last5_case_rows,
    ACTUALIZE_HUILE_NOTE,
    CASE_SIZE_PURCHASE_UNIT_NOTE,
    COUNTSHEET_TRACK_NOTE,
    OPENING_INVENTORY_NOTE, OPENING_INVENTORY_FLOW, OPENING_INVENTORY_COUNT_NOTE,
    OPENING_INVENTORY_VIEW_NOTE, OPENING_INVENTORY_VIEW_HEADERS, opening_inventory_view_rows,
    INVOICE_PHASE_NOTE, INVOICE_MANUAL_NOTE, INVOICE_ENTRY_FLOW,
    INVOICE_SCREEN_HEADERS, invoice_screen_rows,
    INVOICE_LINE_HEADERS, invoice_line_rows,
    INVOICE_SUMMARY_HEADERS, invoice_summary_rows,
    INVOICE_LAB_BOULANGERIE, INVOICE_LAB_DISTRIB_LUNDI,
    EXO_DATE_OPENING, EXO_DATE_CLOSING, EXO_WEEK_NOTE, EXO_WEEK_HEADERS, exo_week_rows, EXO_DAY_TO_DATE, EXO_FIRST_SALES_DAY,
    ORDER_MARDI_HEADERS, order_mardi_rows,
    OPENING_INVENTORY_WIZARD_HEADERS, opening_inventory_wizard_rows,
    OPENING_INVENTORY_QTY_HEADERS, opening_inventory_qty_rows,
    PREPS_NOTE, PREPS_NESTED_NOTE, PREPS_YIELD_NOTES, PREPS_BATCH_UNIT_NOTE, PREPS_ACTUALIZE_NOTE, PREPS_SCREEN_INTRO,
    PREPS_FIELD_GUIDE_HEADERS, prep_field_guide_rows,
    PREPS_CORE_HEADERS, prep_core_rows,
    PREPS_SUMMARY_HEADERS, prep_summary_rows,
    PREPS_INGREDIENT_HEADERS, prep_ingredient_rows,
    PREPS_COST_HINTS, prep_cost_hint_rows,
    PRODUCTS_NOTE, PRODUCTS_NESTED_NOTE, PRODUCTS_TOMATES_NOTE, SALADE_CESAR_NOTE, PRODUCTS_SCREEN_INTRO,
    PRODUCTS_FIELD_GUIDE_HEADERS, product_field_guide_rows,
    PRODUCTS_CORE_HEADERS, product_core_rows,
    PRODUCTS_INGREDIENT_HEADERS, product_ingredient_rows,
    POS_ID_NOTE, POS_ID_CHAIN_NOTE, POS_ID_NOT_NOTE, POS_ID_TILL_TAPE_NOTE, POS_ID_FIELD_HINT,
    POS_ID_MULTI_HEADERS, pos_id_multi_rows, POS_ID_COMPTOIR_HEADERS, pos_id_comptoir_rows,
    publish_docx,
)

OUT = Path(r"E:\OC DOCS\Exercice_Parcours_Lineaire.docx")
REF_EXO = "Exercice_Pratique_Optimum_Control.docx"

STEPS_OVERVIEW = [
    (1, "Comprendre OC (triangle Items → Preps → Products)", "#00–01", "A"),
    (2, "Company Information et magasin (Store)", "Assistant + Store", "A"),
    (3, "Plan comptable (Accounts / GL)", "#02 + #49", "A"),
    (4, "Categories (+ Income Account)", "#02", "A"),
    (5, "Inventory Groups (Sales Cat. + Account)", "#02", "A"),
    (6, "Sales Groups (Sales Cat.)", "#02", "A"),
    (7, "Emplacements de stockage (Storage Locations)", "#02", "A"),
    (8, "Fournisseurs (Suppliers)", "#02", "A"),
    (9, "Unités de mesure (UOM) et conversions", "#02 + ch.12", "A"),
    (10, "Taxes, ajustements et Tax Groups", "#03", "A"),
    (11, "Préférences — Inventory", "Preferences", "A"),
    (12, "Préférences — Purchasing", "Preferences + #03", "A"),
    (13, "Préférences — Sales et Accounting", "Preferences + #49", "A"),
    (14, "Import POS (Preferences → POS)", "#42", "A"),
    (15, "Employés et niveaux d'accès", "#04–05", "A"),
    (16, "CHECKPOINT — configuration terminée", "—", "A"),
    (17, "Créer items 1/3 (pain, bœuf, pommes)", "#06", "B"),
    (18, "Créer items 2/3 (alimentaire)", "#06", "B"),
    (19, "Créer items 3/3 (papier, boissons)", "#06", "B"),
    (20, "Conversions et Case Size Overview", "#08–10", "B"),
    (21, "Item Par Levels", "#11–12", "B"),
    (22, "Créer les preps", "#16", "C"),
    (23, "Recipe Costing", "#18", "C"),
    (24, "Créer les products + POS ID", "#20–21", "C"),
    (25, "Countsheet Setup", "#22", "D"),
    (26, "Custom Sort (shelf-to-sheet)", "#24", "D"),
    (27, "Inventaire d'ouverture", "#25", "D"),
    (28, "Hot List", "#15", "D"),
    (29, "Facture manuelle", "#37", "E"),
    (30, "Créer une commande", "#29–32", "E"),
    (31, "Order Reminders", "#31", "E"),
    (32, "Till Tape — mardi", "#43", "F"),
    (33, "Import CSV — mar→ven", "#42", "F"),
    (34, "Pending Sales", "#44", "F"),
    (35, "Waste", "#45", "F"),
    (36, "Daily Sales", "ch.10", "F"),
    (37, "Inventaire de clôture", "#25", "G"),
    (38, "Usage Summary", "#48", "G"),
    (39, "Recipe Book + Backup", "#19 + #51", "G"),
    (40, "FIN — parcours complet", "—", "G"),
]

def H(doc, t, level=1):
    return doc.add_heading(t, level=level)

def P(doc, t):
    doc.add_paragraph(t)

def bullets(doc, items):
    for i in items:
        doc.add_paragraph(i, style="List Bullet")

def checks(doc, items):
    for i in items:
        doc.add_paragraph(f"☐ {i}", style="List Paragraph")

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

def note(doc, text):
    p = doc.add_paragraph()
    p.add_run("Note — ").bold = True
    p.add_run(text)

def step(doc, n, title, video, prereq, objectif, actions, data_headers=None, data_rows=None, data2_title=None, data2_headers=None, data2_rows=None, data3_title=None, data3_headers=None, data3_rows=None, data4_title=None, data4_headers=None, data4_rows=None, validate=None, legend_headers=None, legend_rows=None):
    H(doc, f"ÉTAPE {n} — {title}", 2)
    p = doc.add_paragraph()
    p.add_run("Tutoriel / menu OC : ").bold = True
    p.add_run(video)
    P(doc, f"Objectif : {objectif}")
    if prereq:
        p2 = doc.add_paragraph()
        p2.add_run("Prérequis : ").bold = True
        p2.add_run(prereq)
    H(doc, "À faire (dans l'ordre — comme la vidéo)", 3)
    for i, a in enumerate(actions, 1):
        pa = doc.add_paragraph()
        pa.add_run(f"{i}. ").bold = True
        pa.add_run(a)
    if legend_headers and legend_rows:
        H(doc, "Signification des Account Types", 3)
        tbl(doc, legend_headers, legend_rows)
    if data_headers and data_rows:
        H(doc, "Valeurs — Le Comptoir du Marché", 3)
        tbl(doc, data_headers, data_rows)
    if data2_headers and data2_rows:
        H(doc, data2_title or "Référence supplémentaire", 3)
        tbl(doc, data2_headers, data2_rows)
    if data3_headers and data3_rows:
        H(doc, data3_title or "Détail complémentaire", 3)
        tbl(doc, data3_headers, data3_rows)
    if data4_headers and data4_rows:
        H(doc, data4_title or "Référence additionnelle", 3)
        tbl(doc, data4_headers, data4_rows)
    H(doc, "Validez avant l'étape suivante", 3)
    nxt = n + 1 if n < 40 else "—"
    checks(doc, validate or [f"Étape {n} OK dans OC Pro.", f"Je passe à l'étape {nxt}."])
    doc.add_paragraph("—" * 36)


def build():
    doc = Document()
    for m in doc.sections:
        m.top_margin = Cm(2)
        m.left_margin = Cm(2.5)
        m.right_margin = Cm(2.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PARCOURS LINÉAIRE\n")
    r.bold = True
    r.font.size = Pt(22)
    t.add_run("Tutoriels + configuration complète\n\n").font.size = Pt(14)
    t.add_run("Le Comptoir du Marché · Port-au-Prince, Haïti · 40 étapes\n\n").font.size = Pt(12)
    t.add_run(f"Devise : {DEVISE}").font.size = Pt(11)
    doc.add_page_break()

    H(doc, "Comment utiliser ce document")
    P(doc, "Parcours strictement linéaire : une étape à la fois, comme les vidéos YouTube Optimum Control. Ne passez à l'étape suivante que lorsque la checklist est cochée.")
    note(doc, NOTE_DEVISE)
    P(doc, f"L'exercice thématique ({REF_EXO}) reste disponible en parallèle — mêmes données, organisation par thème.")
    P(doc, "Deux numérotations : ce document = étapes 1 à 40. L'autre = Parties 0 à 8. "
         "Inventaire d'ouverture = étape 27 ici = Partie 5 dans Exercice_Pratique_Optimum_Control.docx.")
    H(doc, "Phases", 2)
    tbl(doc, ["Phase", "Étapes", "Contenu"], [
        ["A", "1 → 16", "Configuration complète (Store, Setup, Preferences, POS, sécurité)"],
        ["B", "17 → 21", "Items d'inventaire"],
        ["C", "22 → 24", "Preps et Products"],
        ["D", "25 → 28", "Inventaire et comptage"],
        ["E", "29 → 31", "Factures et commandes"],
        ["F", "32 → 36", "Ventes, pertes, Daily Sales"],
        ["G", "37 → 40", "Clôture, rapports, backup"],
    ])
    H(doc, "Tableau de progression (40 étapes)", 2)
    tbl(doc, ["#", "Étape", "Réf.", "☐"], [[str(a), b, c, ""] for a, b, c, _ in STEPS_OVERVIEW])
    doc.add_page_break()

    # ===== PHASE A =====
    H(doc, "PHASE A — Configuration complète (étapes 1 à 16)")
    P(doc, "Reproduisez la vidéo #02 Setup Creation section par section, plus Store, Preferences et Taxes (#03).")

    step(doc, 1, "Comprendre OC", "Vidéos #00 + #01",
         "OC Pro installé.",
         "Comprendre la logique avant toute saisie.",
         ["Regardez #00 et #01.", "Repérez : Settings, Items, Recipe, Inventory, Sales, Purchasing, Reports.",
          "Note : Items = achats | Preps = préparations | Products = menu.",
          "Créez la base Comptoir_Marche_TEST."],
         validate=["Base test prête.", "Je passe à l'étape 2."])

    step(doc, 2, "Company Information et magasin (Store)", "Assistant démarrage + Store",
         "Étape 1.",
         "Renseigner l'identité du restaurant (premier lancement ou File → Company Information).",
         ["Complétez Company Information : nom, adresse, téléphone, courriel.",
          "Mailing Address si différent.",
          f"Store / magasin : Le Comptoir du Marché — {ADRESSE}.",
          f"Devise : {DEVISE}. Semaine de rapport : {SEMAINE_RAPPORT}.",
          "Save."],
         ["Champ", "Valeur"],
         [["Nom", "Le Comptoir du Marché"], ["Adresse", ADRESSE],
          ["Devise", DEVISE], ["Semaine rapport", "Dim → Sam"], ["Pays", "Haïti"]])

    step(doc, 3, "Plan comptable (Accounts)", "Setup → Accounts · Vidéo #02 + #49",
         "Étape 2.",
         "Comptes GL pour factures, taxes et export comptable futur.",
         ["Settings → Setup → Accounts.",
          "Créez tous les comptes du tableau (6 lignes), dont :",
          "• TCA recoverable — Account Type Liability (recommandé), GL 2310 — visible dans Taxes → Account.",
          "• Sales Food / Sales Beverage — Account Type Income (étape 4 Categories).",
          "• Food Cost, Beverage Cost, Paper / Supplies — Account Type CostOfSales (étape 5 Inventory Groups).",
          "Pour chaque ligne : Name + Account Type + GL Account Reference.",
          "Save en haut à droite (comme vidéo #02)."],
         ["Nom (OC)", "GL ref", "Account Type"],
         [[a, b, c] for a, b, c in COMPTES_GL],
         validate=["6 comptes créés dont TCA recoverable (Liability).", "Étape 4."],
         legend_headers=["Account Type", "Signification", "Exemple (restaurant)"],
         legend_rows=[[a, b, c] for a, b, c in ACCOUNT_TYPES_DOC])

    step(doc, 4, "Categories", "Setup → Categories · Vidéo #02",
         "Accounts faits (étape 3).",
         "Créer les catégories et lier chacune à un Income Account.",
         ["Settings → Setup → Categories (écran dédié).",
          "Add : nom de la catégorie (Food, Beverage…) + Income Account.",
          "Income Account = compte Income créé à l'étape 3 (ex. Sales Food).",
          "Save.",
          "Vérification : ces noms apparaîtront dans le dropdown Sales Cat. aux étapes 5 et 6."],
         ["Category", "Income Account", "Usage"],
         [[a, b, c] for a, b, c in CATEGORIES])

    step(doc, 5, "Inventory Groups", "Setup → Inventory Groups · #02",
         "Categories créées (étape 4).",
         "Groupes d'items/preps : lier Sales Cat. + Account.",
         ["Settings → Setup → Inventory Groups (écran séparé de Categories).",
          "Add : Group desc. | Sales Cat. | Account.",
          "Sales Cat. = dropdown : sélectionnez Food, Beverage… (créés étape 4).",
          "Sur une fiche Item : Inventory Group seulement (pas Category). Sales Cat. vient du groupe.",
          "Save."],
         ["Group desc.", "Sales Cat.", "Account"],
         [[a, b, c] for a, b, c in INVENTORY_GROUPS])

    step(doc, 6, "Sales Groups", "Setup → Sales Groups · #02",
         "Inventory Groups faits (étape 5).",
         "Groupes des Products — lier chaque groupe à une Sales Cat.",
         ["Settings → Setup → Sales Groups (écran séparé).",
          "Add : Group desc. + Sales Cat.",
          "Sales Cat. = même dropdown que étape 5 (liste des Categories).",
          "Group desc. = nom aligné sur le POS (peut être plus fin que la catégorie : Entrées, Boissons…).",
          "Save."],
         ["Group desc.", "Sales Cat."],
         [[a, b] for a, b in SALES_GROUPS],
         validate=["Dropdown Sales Cat. affiche Food et Beverage.", "Sales Groups créés.", "Étape 7."])

    step(doc, 7, "Emplacements de stockage (Storage Locations)", "Setup → Storage Locations · #02",
         "Sales Groups faits.",
         "Multi-stocks : Primary + Secondary sur chaque item.",
         ["Setup → Storage Locations → Add (7 emplacements).",
          "Un item peut être assigné à plusieurs emplacements (vidéo #02).",
          "Save."],
         ["Location", "Rôle"],
         [["Cuisine-Froid", "Primary frais"], ["Congélateur", "Primary surgelés"],
          ["Sec", "Primary garde-manger"], ["Ligne-chaude", "Secondary / preps actifs"],
          ["Friterie", "Primary friture"], ["Bar", "Primary boissons"], ["Réserve", "Secondary surplus"]])

    step(doc, 8, "Fournisseurs (Suppliers)", "Setup → Suppliers · #02 + #49",
         "Locations créées.",
         "Vendeurs + Accounting Vendor ID pour export comptable.",
         ["Setup → Suppliers → Add pour chaque fournisseur.",
          "Nom (obligatoire) + Address, City, Province, Postal Code, Country.",
          "Contact, téléphone et courriel si les champs existent sur la fiche.",
          "Accounting Vendor ID = ID du vendeur dans votre logiciel comptable (ex. 576294).",
          "Save."],
         SUPPLIER_HEADERS,
         supplier_table_rows())

    step(doc, 9, "Unités de mesure et conversions", "Setup → Units · Unit Conversions · #02",
         "Suppliers créés.",
         "UOM achat vs recette vs rapport.",
         ["Setup → Units of Measure : vérifiez gramme, kilogramme, ml, litre, chacun, caisse, oz.",
          "Ajoutez si manquant (Add en haut à gauche).",
          "Setup → Unit Conversions (ou Settings) : 1 lb = 453,592 g ; 1 oz liq = 29,5735 ml ; 1 L = 1000 ml.",
          "Save."],
         validate=["UOM standard OK.", "Conversions lb/g et oz/ml actives.", "Étape 10."])

    step(doc, 10, "Taxes, ajustements et Tax Groups", "Setup → Taxes · Tax Groups · #03",
         "UOM configurées.",
         f"Taxes + groupe {TAX_GROUP_CODE} avant toute facture — {TAXES_DOC}",
         ["Setup → Taxes and Adjustments → Add.",
          "Code + Description (ex. TCA — Taxe sur le Chiffre d'Affaires).",
          "Recoverable : Oui (TCA récupérable sur achats).",
          "Type : Percentage — Taux : 10 %.",
          "Account : TCA recoverable (Liability — CostOfSales/Expense/Liability listés, pas Asset/Income).",
          "Apply purchase amount on item : Non (la taxe va au compte TCA, pas au coût item).",
          "Setup → Tax Groups → Add : code " + TAX_GROUP_CODE + ", description Achats soumis à TCA.",
          "Cochez la taxe TCA 10 % dans le groupe → Save.",
          "Sur chaque item (Case Size) : Tax Group = " + TAX_GROUP_CODE + ", Price Includes Tax = décoché.",
          "Save."],
         TAX_HEADERS,
         tax_table_rows(),
         "Tax Group à créer",
         TAX_GROUP_HEADERS,
         tax_group_rows())

    step(doc, 11, "Préférences — Inventory", "Settings → Preferences → Inventory",
         "Taxes créées.",
         "Valorisation du stock, comptage et prep sheets.",
         ["Settings → Preferences → Inventory.",
          "Value Inventory Using → FIFO (voir tableau).",
          "Ask to Summarize on Close → coché.",
          "Require Reason for Inventory Adjustments → coché.",
          "Countsheet Column Display → All Columns.",
          "Inventory Interface → [None].",
          "Default Prep Margin 10 % · Prep Amount Factor 100 %.",
          "Warning Threshold 2 · Critical Threshold 3.",
          "Note : Key Item et Actualize se règlent sur la fiche Item/Prep (étapes items/preps), pas ici.",
          "Save."],
         INVENTORY_PREFS_HEADERS,
         inventory_prefs_rows())

    step(doc, 12, "Préférences — Purchasing", "Settings → Preferences → Purchasing · #03",
         "Prefs Inventory OK.",
         "Factures et commandes.",
         ["Preferences → Purchasing.",
          "Show advanced tax adjustment fields in invoice window → ON (obligatoire vidéo #03).",
          "Require order approval → OFF pour l'instant (ON = vidéo #33 en annexe).",
          "Save."])

    step(doc, 13, "Préférences — Sales et Accounting", "Preferences · #49",
         "Prefs Purchasing OK.",
         "Ventes et export comptable.",
         ["Preferences → Sales : format dates, export comptable si visible.",
          "Preferences → Accounting : système comptable (QuickBooks, Sage, etc.) + dossier export.",
          "Average Hourly Wage si vous saisirez Labour (Daily Sales).",
          "Save."])

    step(doc, 14, "Import POS (Preferences → POS)", "Settings → Preferences → POS · #42",
         "Prefs Sales/Accounting OK.",
         "Préparer l'import CSV des ventes (étape 33).",
         ["Settings → Preferences → POS (pas Setup → Configure POS en v5).",
          "Default POS import folder → parcourir vers exercice_comptoir/ventes_csv/.",
          "POS configuration : sélectionnez ou créez [New Pos Specification].",
          "Engrenage → écran 1 : File Format Csv, Extension .csv, Date Format yyyy-MM-dd, Delimiter virgule, Lines To Skip 1.",
          "Engrenage → écran 2 (Import Specification Fields) : voir tableau ci-dessous (Index = n° de colonne).",
          "Update selling price after import → Never.",
          "Default Category when Creating Product → Food.",
          "Type of sales in reporting → Gross Sales.",
          "Import Sales at Start → décoché.",
          "Save."],
         POS_PREFS_HEADERS,
         pos_prefs_rows(),
         "Mappage Import Specification Fields (Index = n° colonne)",
         POS_SPEC_HEADERS,
         pos_spec_rows())

    step(doc, 15, "Employés et niveaux d'accès", "Security · #04–05",
         "Setup complet.",
         "Comptes utilisateurs.",
         ["Security → Employees → New : Manager (vous) + Employé démo.",
          "Access Levels : Full Access vs Update Not Create.",
          "Reconnectez-vous avec Manager.",
          "Save."])

    step(doc, 16, "CHECKPOINT — configuration terminée",
         "Récap Phase A",
         "Étapes 1–15 faites.",
         "Vérifier que TOUT le Setup est en place avant le premier item.",
         ["Parcourez Settings → Setup : Accounts, Categories, Inv Groups, Sales Groups, Locations, Suppliers, UOM, Taxes — tout rempli ?",
          "Preferences Inventory / Purchasing / Sales / Accounting — Save fait ?",
          "Preferences → POS : dossier CSV + spec colonnes définis ?",
          "2 employés actifs ?"],
         validate=[
             "☑ Store / Company Information",
             "☑ Accounts (Name + Type + GL)",
             "☑ Categories + Inventory Groups + Sales Groups",
             "☑ Storage Locations (7)",
             "☑ Suppliers (4) + Accounting Vendor ID",
             "☑ UOM + Conversions",
             "☑ Taxes + Tax Group TCA-ACHAT + Pref Purchasing (advanced tax ON)",
             "☑ Preferences Inventory / Sales / Accounting",
             "☑ Preferences → POS (dossier + spec CSV)",
             "☑ Employees — JE PASSE À LA PHASE B (étape 17)",
         ])

    doc.add_page_break()
    H(doc, "PHASE B — Items (étapes 17 à 21)")

    step(doc, 17, "Créer items 1/3", "#06 Creating Items",
         "CHECKPOINT 16 validé.",
         "Reproduire #06 : 3 premiers items — saisie complète (Core + Case Size + Tax Group).",
         ["Inventory → Items → New.",
          "Core (gauche) : voir tableau 1 — Description, Inventory Group, Main Location, How is it used?, Reporting/Ingredient UOM, Key Item, Actualize.",
          "Case Size (droite) : Supplier ; Purchase Price + liste Prix pour (case, bag, each…) ; Split puis Pack (tableau 2).",
          "Tax Group, Yield %, Price Includes Tax = Non.",
          f"Tax Group obligatoire : {TAX_GROUP_CODE} (créé étape 10).",
          "Save → onglet Locations si emplacement secondaire (pommes : option Friterie en secondaire).",
          "Ne jamais dupliquer un item : 2e fournisseur = Add (case size)."],
         ITEMS_FIRST3_CORE_HEADERS,
         items_first3_core_rows(),
         "Case Size — 3 premiers items",
         ITEMS_FIRST3_CASE_HEADERS,
         items_first3_case_rows(),
         validate=[
             f"3 items créés avec Tax Group = {TAX_GROUP_CODE}.",
             "Coût unitaire OC vérifié (pain 0,20 $ ; bœuf 0,0084 $/g ; pommes ≈ 0,0011 $/g).",
             "Étape 18.",
         ])

    step(doc, 18, "Créer items 2/3 (alimentaire)", "#06 suite",
         "3 items créés.",
         "13 items alimentaires restants — même méthode Core + Case Size que l'étape 17.",
         [ITEM_MULTI_CASESIZE_NOTE,
          ITEMS_REST13_NOTE,
          ACTUALIZE_HUILE_NOTE,
          ITEM_MULTI_CASESIZE_EXAMPLE_NOTE,
          "Inventory → Items → New pour chaque ligne des tableaux.",
          "Core (gauche) : tableau 1. Case Size (droite) : tableau 2 — Supplier, prix, Split/Pack, Yield %, Tax Group.",
          f"Tax Group = {TAX_GROUP_CODE}, Price Includes Tax = Non sur chaque case size.",
          "Save → onglet Locations : ketchup, mayo, bacon → ajouter Ligne chaude en secondaire.",
          "Exercice labo multi-fournisseur — voir tableau 3 (Bacon, 2e case size)."],
         ITEMS_DETAIL_CORE_HEADERS,
         items_rest13_core_rows(),
         "Case Size — 13 items alimentaires",
         ITEMS_DETAIL_CASE_HEADERS,
         items_rest13_case_rows(),
         "Exercice labo — Bacon tranché, 2e case size (Add)",
         ITEM_MULTI_CASESIZE_EXAMPLE_HEADERS,
         items_multicasize_example_rows(),
         validate=[
             "13 items alimentaires créés (total 16 avec étape 17).",
             f"Chaque case size a Tax Group = {TAX_GROUP_CODE}.",
             "Huile friture : Actualize = Oui. Sel : Key Item = Non (option).",
             "Bacon : 2 case sizes (Distrib. Caraïbes + Boulangerie Pétion) — View All = 2 lignes, 1 seul item.",
             "Étape 19.",
         ])

    step(doc, 19, "Créer items 3/3 (papier, boissons)", "#06 fin",
         "16 items alimentaires créés.",
         "5 derniers items — boissons et papier. 21 items au total.",
         [ITEMS_LAST5_NOTE,
          "Même structure Core + Case Size que les étapes 17–18.",
          f"Tax Group = {TAX_GROUP_CODE} sur chaque case size.",
          "Comptez : 21 items actifs au total."],
         ITEMS_DETAIL_CORE_HEADERS,
         items_last5_core_rows(),
         "Case Size — boissons et papier",
         ITEMS_DETAIL_CASE_HEADERS,
         items_last5_case_rows(),
         validate=[
             "5 items boissons/papier créés.",
             "21 items actifs au total.",
             "Étape 20.",
         ])

    step(doc, 20, "Conversions et Case Size Overview", "#08–10",
         "21 items.", "Vérifier case sizes.", ["Case Size Overview sur bœuf et ketchup.", "0,0084 $US/g bœuf OK ?"])

    step(doc, 21, "Item Par Levels", "#11–12",
         "Items validés.", "Par levels.", ["Item Par Levels : pain 12/36, bœuf 5000/15000 g, cola 24/72.", "Save."])

    doc.add_page_break()
    H(doc, "PHASE C — Recettes (22–24)")

    step(doc, 22, "Créer les preps", "#16 Creating Preps",
         "Par levels OK (étape 21).",
         "5 preps — chaque champ de l'écran New Prep renseigné (voir tableaux).",
         [PREPS_SCREEN_INTRO,
          PREPS_NOTE,
          PREPS_NESTED_NOTE,
          PREPS_YIELD_NOTES,
          PREPS_BATCH_UNIT_NOTE,
          PREPS_ACTUALIZE_NOTE,
          "Batch Uom = Batch ; Batch Yield = 1 batch ; Qty recette = poids/volume produit (4000 g, 950 ml…).",
          "Recipe → Preps → New — reproduire chaque prep (ordre : batch frites avant portion frites).",
          "Core : Description, How is it used?, Batch Yield, Recipe Unit — tableau « Valeurs par prep ».",
          "Inventory : Count sheet + Main Location. Prep Sheet : Include, Station, Shelf Life.",
          "Ingrédients : Add Item ou Add Prep (icônes en haut du tableau droit).",
          "Save → vérifier Actual Cost Per (tableau coûts) à l'étape 23."],
         PREPS_FIELD_GUIDE_HEADERS,
         prep_field_guide_rows(),
         "Valeurs par prep — Le Comptoir du Marché",
         PREPS_CORE_HEADERS,
         prep_core_rows(),
         "Ingrédients — Items et Preps",
         PREPS_INGREDIENT_HEADERS,
         prep_ingredient_rows(),
         "Coûts attendus (vérification étape 23)",
         ["Prep", "Coût unitaire OC", "Calcul"],
         prep_cost_hint_rows(),
         validate=[
             "5 preps : Core, Inventory, Prep Sheet et Ingrédients complets.",
             "Portion frites : Add Prep → batch frites en ingrédient.",
             "Étape 23.",
         ])

    step(doc, 23, "Recipe Costing", "#18",
         "Preps créés.", "Vérifier coûts.", ["Notez cost/unit sur chaque prep."])

    step(doc, 24, "Créer les products + POS ID", "#20–21",
         "Preps OK.", "7 products liés POS — champs New Product complets.",
         [POS_ID_NOTE,
          POS_ID_CHAIN_NOTE,
          POS_ID_NOT_NOTE,
          POS_ID_TILL_TAPE_NOTE,
          POS_ID_FIELD_HINT,
          PRODUCTS_SCREEN_INTRO,
          PRODUCTS_NOTE,
          PRODUCTS_NESTED_NOTE,
          PRODUCTS_TOMATES_NOTE,
          "Sales Group obligatoire (Entrées / Plats, Accompagnements, Boissons).",
          "POS ID# 101–502 · Selling Price menu · ingrédients Add Item / Prep / Product.",
          "Burger bacon : Add Product Burger classique + 2 bacon.",
          "Salade César 301 : 4 Items — laitue 150 g, dressing 45 ml, parmesan 15 g, croûtons 20 g (pas de poulet — voir note ci-dessous).",
          SALADE_CESAR_NOTE,
          "Bol chili 401 : Prep Chili — batch 400 g + fourchette. Cola/Eau : 1 item each, Sales Group Boissons.",
          "Save → noter Cost Percent (food cost %) sur chaque product."],
         POS_ID_MULTI_HEADERS,
         pos_id_multi_rows(),
         "Table PLU — Le Comptoir du Marché",
         POS_ID_COMPTOIR_HEADERS,
         pos_id_comptoir_rows(),
         "Champs New Product",
         PRODUCTS_FIELD_GUIDE_HEADERS,
         product_field_guide_rows(),
         "Valeurs par product — Core et Pricing",
         PRODUCTS_CORE_HEADERS,
         product_core_rows(),
         validate=[
             "7 products : Sales Group + POS ID + ingrédients complets (voir tableau ingrédients dans l'exercice complet).",
             "Table PLU 101–502 comprise (1 POS ID = 1 Product).",
             "Cost Percent noté — étape 25.",
         ])

    doc.add_page_break()
    H(doc, "PHASE D — Inventaire (25–28)")

    step(doc, 25, "Countsheet Setup", "#22", "Products OK.", "Feuilles comptage.",
         ["Count Inventory → Countsheet Setup (vidéo #22).",
          "Track Inventory : coché = item suivi (qty on hand) — c'est ici en v5, pas sur la fiche Item.",
          "Should Count : default case size (icône panier) toujours sur feuille ; cocher autres case sizes si besoin.",
          "Save."])

    step(doc, 26, "Custom Sort", "#24", "Countsheet OK.", "Shelf-to-sheet.",
         ["Customize Sort : Froid, Sec, Congélateur → Save."])

    step(doc, 27, "Inventaire d'ouverture", "#25", "Sort OK.", "1er inventaire — saisir cs (items) ou batch (preps).",
         [OPENING_INVENTORY_NOTE,
          OPENING_INVENTORY_COUNT_NOTE,
          OPENING_INVENTORY_VIEW_NOTE,
          f"Count Inventory → New → Create Inventory : date **{EXO_DATE_OPENING}** (lundi soir), All Items, Multiple sheets décoché → Finish.",
          "Tri Location → saisir le tableau « Saisie inventaire d'ouverture » ci-dessous.",
          "Summarize Count → vérifier totaux → Finalize (pas de factures avant cette étape)."],
         OPENING_INVENTORY_QTY_HEADERS,
         opening_inventory_qty_rows(),
         data2_title="Assistant Create Inventory (avant la feuille)",
         data2_headers=OPENING_INVENTORY_WIZARD_HEADERS,
         data2_rows=opening_inventory_wizard_rows(),
         data3_title="Où voir l'inventaire et Qty on Hand (OC v5)",
         data3_headers=OPENING_INVENTORY_VIEW_HEADERS,
         data3_rows=opening_inventory_view_rows(),
         validate=[
             f"Inventaire d'ouverture Finalize (date = {EXO_DATE_OPENING}, lundi).",
             "Items : Qty on Hand sur Inventory → Items (ex. Pain = 24).",
             "Preps : pas sur fiche Prep — Count Inventory ou Prep Sheet On Hand.",
             "Je passe à l'étape 28.",
         ])

    step(doc, 28, "Hot List", "#15", "Ouverture Finalize.", "Comptage partiel.",
         ["Hot List « Frais » : laitue, tomates, parmesan, chili prep."])

    doc.add_page_break()
    H(doc, "PHASE E — Achats (29–31)")
    P(doc, INVOICE_PHASE_NOTE)
    P(doc, INVOICE_MANUAL_NOTE)
    P(doc, EXO_WEEK_NOTE)
    tbl(doc, EXO_WEEK_HEADERS, exo_week_rows())

    step(doc, 29, "Facture manuelle — labo + semaine", "#37", "Ouverture Finalize.", "Saisir 5 factures ligne par ligne.",
         INVOICE_ENTRY_FLOW + [
          "Facture 1 (labo) — reproduire le bon de commande papier ci-dessous dans OC.",
          INVOICE_LAB_BOULANGERIE,
          "Facture 2 (mardi) — Distrib. Caraïbes : " + INVOICE_LAB_DISTRIB_LUNDI,
          "Factures 3–5 — mercredi / jeudi / vendredi : voir tableaux « Lignes » et « Totaux ».",
         ],
         INVOICE_SCREEN_HEADERS,
         invoice_screen_rows(),
         data2_title="Lignes de facture — semaine complète (copier dans OC)",
         data2_headers=INVOICE_LINE_HEADERS,
         data2_rows=invoice_line_rows(),
         data3_title="Totaux par facture (Invoice Total = Total TTC)",
         data3_headers=INVOICE_SUMMARY_HEADERS,
         data3_rows=invoice_summary_rows(),
         validate=[
             "5 factures Save — Account Balance = 0 sur chacune.",
             "Qty on Hand augmentée (ex. Pain : 24 ouverture + 24 facture mardi).",
             "Popup Price Variance lu — OK si 1er achat.",
             "Je passe à l'étape 30.",
         ])

    step(doc, 30, "Créer une commande", "#32", "Facture mardi faite.", "Commande fournisseur.",
         ["Purchasing → Orders → New (New Order Worksheet).",
          f"Supplier = {NOM_DISTRIB}.",
          "Ajouter les items du tableau (double-clic ou glisser).",
          "Round to Case si proposé · Set Approved si Order Approval actif · Save.",
          "Ne pas exporter EDI pour l'exercice seul."],
         ORDER_MARDI_HEADERS,
         order_mardi_rows(),
         validate=["1 commande Save.", "Étape 31."])

    step(doc, 31, "Order Reminders", "#31", "Par levels configurés (étape 21).", "Dashboard réappro.",
         ["Dashboard → tuile Order Reminder (items sous le minimum).",
          "Create Order → Round to Case → Save par fournisseur.",
          "Option : cocher Include items below par level lors d'une commande."],
         validate=["Tuile Order Reminder consultée.", "Phase F — ventes."])

    doc.add_page_break()
    H(doc, "PHASE F — Ventes (32–36)")
    P(doc, TILL_TAPE_NOTE)
    P(doc, TILL_TAPE_TWO_STEP_NOTE)

    step(doc, 32, "Till Tape — mardi", "#43", "Products POS OK.", f"Liste Till Tape puis vente {EXO_DAY_TO_DATE[EXO_FIRST_SALES_DAY]}.",
         TILL_TAPE_LIST_FLOW + TILL_TAPE_SALE_FLOW + [TILL_TAPE_POS_NOTE],
         TILL_TAPE_QTY_HEADERS,
         till_tape_mardi_rows(),
         validate=TILL_TAPE_VALIDATE + ["Je passe à l'étape 33."])

    step(doc, 33, "Import CSV mer→sam", "#42", "Till Tape mar. fait.", "4 jours CSV — Comptoir CSV actif.",
         ["Garder Comptoir CSV dans Preferences → POS (ne pas le désélectionner).",
          CSV_IMPORT_NOTE,
          "Sales → New → Import From File → Browse.",
          "Import ventes_mercredi_2026-07-01.csv … ventes_samedi_2026-07-04.csv (5 colonnes, Date en col. 1).",
          "Ne pas utiliser ventes_minimal_*.csv avec cette spec.",
          f"Dates : mer. {EXO_DAY_TO_DATE['Mercredi']}, jeu. {EXO_DAY_TO_DATE['Jeudi']}, ven. {EXO_DAY_TO_DATE['Vendredi']}, sam. {EXO_DAY_TO_DATE['Samedi']}."],
         CSV_IMPORT_HEADERS,
         csv_import_map_rows(),
         validate=["4 jours CSV importés ou ventes_semaine_comptoir.csv.", "Pending Sales traités.", "Étape 34."])

    step(doc, 34, "Pending Sales", "#44", "Imports mer→sam faits.", "0 Unlinked — tout Valid.",
         [PENDING_SALES_NOTE, POS_ID_CHAIN_NOTE, PENDING_SALES_UNLINK_NOTE, PENDING_SALES_COMPTOIR_NOTE] + PENDING_SALES_FLOW,
         PENDING_SALES_STATES_HEADERS,
         pending_sales_state_rows(),
         validate=PENDING_SALES_VALIDATE + ["Étape 35."])

    step(doc, 35, "Waste", "#45", "Pending Sales OK.", "4 pertes saisies.",
         [WASTE_NOTE] + WASTE_FLOW,
         WASTE_TABLE_HEADERS,
         waste_rows(),
         validate=WASTE_VALIDATE + ["Étape 36."])

    step(doc, 36, "Daily Sales", "ch.10", "Waste OK.", "5 jours — totaux + champs ops.",
         [DAILY_SALES_NOTE, DAILY_SALES_SCREEN_NOTE, DAILY_SALES_GROSS_NOTE, DAILY_SALES_DISTRIB_NOTE] + DAILY_SALES_FLOW,
         DAILY_SALES_TABLE_HEADERS,
         daily_sales_rows(),
         "Lexique champs Daily Sales",
         DAILY_SALES_FIELDS_HEADERS,
         daily_sales_field_rows(),
         "Valeurs opérationnelles par jour (— = laisser 0)",
         DAILY_SALES_OPS_HEADERS,
         daily_sales_ops_rows(),
         validate=DAILY_SALES_VALIDATE + ["Phase G — étape 37."])

    doc.add_page_break()
    H(doc, "PHASE G — Clôture (37–40)")

    step(doc, 37, "Inventaire de clôture", "#25", "Semaine complète.", f"2e inventaire (samedi {EXO_DATE_CLOSING}).",
         [CLOSING_INVENTORY_NOTE] + CLOSING_INVENTORY_FLOW,
         CLOSING_INVENTORY_QTY_HEADERS,
         closing_inventory_qty_rows(),
         validate=CLOSING_INVENTORY_VALIDATE + ["Étape 38."])

    step(doc, 38, "Usage Summary", "#48", "Clôture save.", "Actual vs Ideal.",
         [USAGE_SUMMARY_NOTE, USAGE_SUMMARY_FORMULA] + USAGE_SUMMARY_FLOW + USAGE_SUMMARY_QUESTIONS,
         validate=USAGE_SUMMARY_VALIDATE + ["Étape 39."])

    step(doc, 39, "Recipe Book + Backup", "#19 + #51", "Usage Summary OK.", "PDF + sauvegarde.",
         [BACKUP_NOTE] + BACKUP_FLOW)

    step(doc, 40, "FIN — parcours complet", "Récap",
         "Étapes 1–39.", "Restaurant opérationnel.",
         ["Cochez tout.", "Consultez Exercice_Corrige_Detaille.docx maintenant."],
         validate=["PARCOURS TERMINÉ ✓"])

    doc.add_page_break()
    H(doc, "Annexe — Setup : carte des menus OC")
    tbl(doc, ["Paramètre", "Menu OC"], [
        ["Store / Company", "File → Company Information · Store"],
        ["Accounts (Name + Type + GL)", "Settings → Setup → Accounts"],
        ["Categories", "Settings → Setup → Categories"],
        ["Inventory Groups", "Settings → Setup → Inventory Groups"],
        ["Sales Groups", "Settings → Setup → Sales Groups"],
        ["Inventory Groups", "Settings → Setup → Inventory Groups"],
        ["Sales Groups", "Settings → Setup → Sales Groups"],
        ["Storage Locations", "Settings → Setup → Storage Locations"],
        ["Suppliers", "Settings → Setup → Suppliers"],
        ["Units of Measure", "Settings → Setup → Units of Measure"],
        ["Unit Conversions", "Settings → Setup → Unit Conversions"],
        ["Taxes & Adjustments", "Settings → Setup → Taxes and Adjustments"],
        ["Tax Groups", "Settings → Setup → Tax Groups"],
        ["Preferences Inventory", "Settings → Preferences → Inventory"],
        ["Preferences Purchasing", "Settings → Preferences → Purchasing"],
        ["Preferences Sales", "Settings → Preferences → Sales"],
        ["Preferences Accounting", "Settings → Preferences → Accounting"],
        ["Import POS (Preferences → POS)", "Settings → Preferences → POS"],
        ["Employees", "Security → Employees"],
        ["Access Levels", "Security → Access Levels"],
    ])
    note(doc, f"Exercice thématique détaillé (multi-stocks, défis bonus) : {REF_EXO}")

    publish_docx(doc, OUT)

if __name__ == "__main__":
    build()
