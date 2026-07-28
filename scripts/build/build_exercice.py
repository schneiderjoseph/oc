#!/usr/bin/env python3
"""Génère l'exercice pratique complet Optimum Control."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from exercice_locale import (
    ADRESSE, DEVISE, DEVISE_COURTE, NOTE_DEVISE, SEMAINE_RAPPORT, TAXES_DOC, TAXES_NOTE, TAX_HEADERS, tax_table_rows,
    TAX_GROUPS_NOTE, TAX_GROUP_HEADERS, tax_group_rows, TAX_GROUP_CODE,
    FOURNISSEURS, NOM_DISTRIB, NOM_BOULANGERIE, NOM_EMBALLAGES, NOM_BOISSONS, COMPTES_GL,
    ACCOUNT_TYPES_DOC, CATEGORIES, INVENTORY_GROUPS, SALES_GROUPS, SETUP_CATEGORIES_NOTE,
    SUPPLIERS_NOTE, SUPPLIER_HEADERS, supplier_table_rows,
    INVENTORY_PREFS_NOTE, INVENTORY_PREFS_HEADERS, inventory_prefs_rows,
    POS_PREFS_NOTE, POS_PREFS_HEADERS, pos_prefs_rows,
    POS_SPEC_HEADERS, pos_spec_rows,
    TILL_TAPE_NOTE, TILL_TAPE_TWO_STEP_NOTE, TILL_TAPE_POS_NOTE,
    TILL_TAPE_LIST_FLOW, TILL_TAPE_SALE_FLOW, TILL_TAPE_VALIDATE,
    TILL_TAPE_QTY_HEADERS, till_tape_mardi_rows, DAILY_SALES_GROSS_NOTE,
    PENDING_SALES_NOTE, PENDING_SALES_UNLINK_NOTE, PENDING_SALES_FLOW, PENDING_SALES_COMPTOIR_NOTE,
    PENDING_SALES_VALIDATE, PENDING_SALES_STATES_HEADERS, pending_sales_state_rows,
    WASTE_NOTE, WASTE_FLOW, WASTE_VALIDATE, WASTE_TABLE_HEADERS, waste_rows,
    DAILY_SALES_NOTE, DAILY_SALES_SCREEN_NOTE, DAILY_SALES_GROSS_NOTE, DAILY_SALES_DISTRIB_NOTE,
    DAILY_SALES_FLOW, DAILY_SALES_VALIDATE,
    DAILY_SALES_FIELDS_HEADERS, daily_sales_field_rows,
    DAILY_SALES_OPS_HEADERS, daily_sales_ops_rows,
    DAILY_SALES_TABLE_HEADERS, daily_sales_rows,
    CLOSING_INVENTORY_NOTE, CLOSING_INVENTORY_FLOW, CLOSING_INVENTORY_VALIDATE,
    CLOSING_INVENTORY_QTY_HEADERS, closing_inventory_qty_rows,
    USAGE_SUMMARY_NOTE, USAGE_SUMMARY_FORMULA, USAGE_SUMMARY_FLOW, USAGE_SUMMARY_QUESTIONS, USAGE_SUMMARY_VALIDATE,
    BACKUP_NOTE, BACKUP_FLOW,
    CSV_IMPORT_NOTE, CSV_IMPORT_HEADERS, csv_import_map_rows,
    ITEM_CREATE_NOTE, ITEM_CREATE_HEADERS, item_create_rows,
    ITEMS_FIRST3_NOTE, ITEMS_FIRST3_CORE_HEADERS, items_first3_core_rows,
    ITEMS_FIRST3_CASE_HEADERS, items_first3_case_rows,
    ITEM_MULTI_CASESIZE_NOTE,
    ITEM_MULTI_CASESIZE_EXAMPLE_NOTE, ITEM_MULTI_CASESIZE_EXAMPLE_HEADERS,
    items_multicasize_example_rows,
    ITEMS_REST13_NOTE, ITEMS_DETAIL_CORE_HEADERS, items_rest13_core_rows,
    ITEMS_DETAIL_CASE_HEADERS, items_rest13_case_rows,
    ITEMS_LAST5_NOTE, items_last5_core_rows, items_last5_case_rows,
    ACTUALIZE_HUILE_NOTE,
    CASE_SIZE_PURCHASE_UNIT_NOTE,
    COUNTSHEET_TRACK_NOTE,
    OPENING_INVENTORY_NOTE, OPENING_INVENTORY_FLOW, OPENING_INVENTORY_COUNT_NOTE,
    OPENING_INVENTORY_VIEW_NOTE, OPENING_INVENTORY_VIEW_HEADERS, opening_inventory_view_rows,
    INVOICE_PHASE_NOTE, INVOICE_MANUAL_NOTE, INVOICE_ENTRY_FLOW,
    INVOICE_SCREEN_HEADERS, invoice_screen_rows,
    INVOICE_LINE_HEADERS, invoice_line_rows,
    INVOICE_SUMMARY_HEADERS, invoice_summary_rows,
    INVOICE_LAB_BOULANGERIE, INVOICE_LAB_DISTRIB_LUNDI,
    EXO_WEEK_NOTE, EXO_WEEK_HEADERS, exo_week_rows, EXO_DATE_OPENING, EXO_DATE_CLOSING, EXO_FIRST_SALES_DAY, EXO_DAY_TO_DATE,
    ORDER_MARDI_HEADERS, order_mardi_rows,
    OPENING_INVENTORY_WIZARD_HEADERS, opening_inventory_wizard_rows,
    OPENING_INVENTORY_QTY_HEADERS, opening_inventory_qty_rows,
    PREPS_NOTE, PREPS_NESTED_NOTE, PREPS_YIELD_NOTES, PREPS_BATCH_UNIT_NOTE, PREPS_ACTUALIZE_NOTE, PREPS_SCREEN_INTRO,
    PREPS_FIELD_GUIDE_HEADERS, prep_field_guide_rows,
    PREPS_CORE_HEADERS, prep_core_rows,
    PREPS_SUMMARY_HEADERS, prep_summary_rows,
    PREPS_INGREDIENT_HEADERS, prep_ingredient_rows,
    PREPS_COST_HINTS, prep_cost_hint_rows,
    PRODUCTS_NOTE, PRODUCTS_NESTED_NOTE, PRODUCTS_TOMATES_NOTE, SALADE_CESAR_NOTE, PRODUCTS_SCREEN_INTRO,
    PRODUCTS_FIELD_GUIDE_HEADERS, product_field_guide_rows,
    PRODUCTS_CORE_HEADERS, product_core_rows,
    PRODUCTS_INGREDIENT_HEADERS, product_ingredient_rows,
    POS_ID_NOTE, POS_ID_CHAIN_NOTE, POS_ID_NOT_NOTE, POS_ID_TILL_TAPE_NOTE, POS_ID_FIELD_HINT,
    POS_ID_MULTI_HEADERS, pos_id_multi_rows, POS_ID_COMPTOIR_HEADERS, pos_id_comptoir_rows,
    publish_docx,
)

OUT = Path(r"E:\OC DOCS\Exercice_Pratique_Optimum_Control.docx")
SUPP_DIR = Path(r"E:\OC DOCS\exercice_comptoir")

def add_title(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

def add_checklist(doc, items):
    for item in items:
        doc.add_paragraph(f"☐ {item}", style="List Paragraph")

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()
    return t

def add_note(doc, text):
    p = doc.add_paragraph()
    r1 = p.add_run("Note — ")
    r1.bold = True
    p.add_run(text)

def add_step(doc, n, text):
    p = doc.add_paragraph()
    r = p.add_run(f"Étape {n}. ")
    r.bold = True
    p.add_run(text)

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # ---- PAGE DE GARDE ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("EXERCICE PRATIQUE\n")
    r.bold = True
    r.font.size = Pt(22)
    t.add_run("Optimum Control — Parcours complet\n").font.size = Pt(16)
    t.add_run("\nGérer un restaurant de A à Z\n").font.size = Pt(14)
    t.add_run("\nScénario : Le Comptoir du Marché — Port-au-Prince, Haïti\n").font.size = Pt(12)
    t.add_run(f"\nDevise : {DEVISE}\n").font.size = Pt(11)
    doc.add_page_break()

    # ---- INTRO ----
    add_title(doc, "Comment utiliser cet exercice")
    add_p(doc, "Cet exercice simule la mise en service complète d'un petit restaurant fictif, Le Comptoir du Marché, dans Optimum Control. Vous configurez le système, créez items, preps et products, puis enchaînez une semaine type : achats, ventes, pertes, inventaire et rapports.")
    add_p(doc, "Durée estimée : 5 à 8 heures (répartissable sur plusieurs sessions).")
    add_note(doc, "Utilisez une base de données de test ou une sauvegarde vierge. Ne travaillez pas sur votre base de production.")
    add_note(doc, NOTE_DEVISE)
    add_p(doc, "Version tutoriel étape par étape (Setup complet, vidéos dans l'ordre) : Exercice_Parcours_Lineaire.docx — 40 étapes, même restaurant. Les deux documents sont complémentaires.")
    add_p(doc, "Équivalence : inventaire d'ouverture = Partie 5 (ce document) = étape 27 du Parcours Linéaire.")
    add_p(doc, "Fichiers complémentaires dans exercice_comptoir/ : ventes CSV, README import, corrigé détaillé.")
    add_title(doc, "Votre mission", 2)
    add_bullets(doc, [
        "Configurer le restaurant et ses fournisseurs.",
        "Créer tous les items d'inventaire avec les bonnes unités et rendements.",
        "Construire les preps (préparations) et les products (articles du menu).",
        "Saisir l'inventaire d'ouverture.",
        "Simuler 5 jours d'exploitation (factures, ventes, pertes).",
        "Faire un inventaire de clôture et analyser le Usage Summary.",
        "Gérer plusieurs emplacements de stockage et des transferts entre stocks.",
        "Interpréter les écarts et corriger une erreur volontaire.",
    ])
    add_title(doc, "Le restaurant — contexte", 2)
    add_p(doc, "Le Comptoir du Marché est un lunch counter urbain (60 couverts/jour en semaine). Menu simple :")
    add_bullets(doc, [
        "Burger classique, Burger bacon, Frites maison",
        "Salade César",
        "Soupe du jour (chili)",
        "Boissons gazeuses 355 ml",
    ])
    add_p(doc, "Objectif food cost cible : 32 % sur la nourriture. Vous travaillez seul(e) avec un comptoir POS fictif (numéros PLU fournis).")
    doc.add_page_break()

    # ---- PARTIE 0 ----
    add_title(doc, "Partie 0 — Avant de commencer")
    add_checklist(doc, [
        "OC Pro installé et activé sur votre poste.",
        "Base de données de test créée ou base vierge restaurée.",
        "Document Comprendre_Optimum_Control_V2_PLAT.docx ouvert en référence.",
        "Bloc-notes pour noter vos Consolidation ID / POS ID si besoin.",
    ])
    add_title(doc, "Triangle à garder en tête", 2)
    add_p(doc, "Items (achats) → Preps (transformations) → Products (ventes). Chaque vente POS liée consomme du stock théorique ; les factures augmentent le stock ; l'inventaire physique révèle les écarts.")
    doc.add_page_break()

    # ---- PARTIE 1 ----
    add_title(doc, "Partie 1 — Configuration initiale complète (Jour 0)")
    add_p(doc, "Cette partie reprend tous les paramètres Setup et Preferences avant de créer le moindre item. Même contenu que les étapes 2–16 du Parcours Linéaire, organisé par thème.")

    add_title(doc, "1.1 Company Information et Store", 2)
    add_step(doc, 1, "File → Company Information (ou assistant au premier lancement) :")
    add_table(doc, ["Champ", "Valeur"], [
        ["Nom", "Le Comptoir du Marché"],
        ["Adresse", ADRESSE],
        ["Devise", DEVISE],
        ["Semaine de rapport", SEMAINE_RAPPORT],
        ["Pays", "Haïti"],
    ])

    add_title(doc, "1.2 Plan comptable (Accounts / GL)", 2)
    add_p(doc, "Settings → Setup → Accounts. Pour chaque compte : Name, Account Type, puis GL Account Reference pour l'export comptable (vidéo #49).")
    add_p(doc, "Important : créez TCA recoverable (type Liability recommandé). Le dropdown Taxes → Account affiche CostOfSales, Expense et Liability — pas Asset ni Income.")
    add_table(doc, ["Account Type", "Signification", "Exemple (restaurant)"], ACCOUNT_TYPES_DOC)
    add_table(doc, ["Nom (OC)", "GL ref", "Account Type"], COMPTES_GL)

    add_title(doc, "1.3 Categories, Inventory Groups et Sales Groups", 2)
    add_p(doc, f"Settings → Setup. {SETUP_CATEGORIES_NOTE}")
    add_p(doc, "Ordre : Accounts → Categories → Inventory Groups → Sales Groups.")
    add_title(doc, "1.3.1 Categories (création + Income Account)", 3)
    add_p(doc, "Setup → Categories. Ici seulement on crée les catégories. Chaque ligne : nom + Income Account (compte Income de l'étape 1.2).")
    add_p(doc, "Ces noms réapparaîtront dans le menu déroulant Sales Cat. / Sales Category sur Inventory Groups et Sales Groups.")
    add_table(doc, ["Category", "Income Account", "Usage"], CATEGORIES)
    add_title(doc, "1.3.2 Inventory Groups", 3)
    add_p(doc, "Setup → Inventory Groups — écran séparé. Colonnes : Group desc. | Sales Cat. | Account.")
    add_p(doc, "Sales Cat. = menu déroulant : choisissez une catégorie créée à l'étape 1.3.1. Account = compte CostOfSales pour les achats (factures).")
    add_p(doc, "Sur une fiche Item : on choisit un Inventory Group (pas une Category). La Sales Cat. est héritée du groupe.")
    add_p(doc, "Ex. : boîtes et serviettes → Inventory Group Paper → Sales Cat. Food + Account Paper / Supplies. Pas de Category Paper.")
    add_table(doc, ["Group desc.", "Sales Cat.", "Account"], INVENTORY_GROUPS)
    add_title(doc, "1.3.3 Sales Groups", 3)
    add_p(doc, "Setup → Sales Groups — écran séparé. Colonnes : Group desc. | Sales Cat.")
    add_p(doc, "Sales Cat. = même menu déroulant (Categories). Group desc. = nom du groupe POS (ex. Entrées, Boissons) — peut différer du nom de la catégorie.")
    add_table(doc, ["Group desc.", "Sales Cat."], SALES_GROUPS)

    add_title(doc, "1.4 Taxes, ajustements et Tax Groups", 2)
    add_p(doc, TAXES_NOTE)
    add_p(doc, "Prérequis : compte TCA recoverable (Liability, GL 2310) créé en partie 1.2.")
    add_p(doc, TAXES_DOC)
    add_table(doc, TAX_HEADERS, tax_table_rows())
    add_title(doc, "1.4.1 Tax Groups (obligatoire)", 3)
    add_p(doc, TAX_GROUPS_NOTE)
    add_table(doc, TAX_GROUP_HEADERS, tax_group_rows())

    add_title(doc, "1.5 Emplacements de stockage (Storage Locations)", 2)
    add_p(doc, "Setup → Storage Locations. Primary + Secondary par item.")
    add_table(doc, ["Emplacement", "Rôle", "Exemples"], [
        ["Cuisine — Froid", "Primary frais", "Laitue, tomates, mayo"],
        ["Congélateur", "Primary surgelés", "Bœuf, bacon"],
        ["Sec", "Primary garde-manger", "Pains, ketchup, haricots"],
        ["Ligne chaude", "Secondary", "Chili, sauce, preps actifs"],
        ["Friterie", "Primary friture", "Huile, pommes"],
        ["Bar", "Primary boissons", "Cola, eau"],
        ["Réserve", "Secondary surplus", "Stock backup, papier vrac"],
    ])

    add_title(doc, "1.6 Fournisseurs (Suppliers)", 2)
    add_p(doc, SUPPLIERS_NOTE)
    add_table(doc, SUPPLIER_HEADERS, supplier_table_rows())

    add_title(doc, "1.7 Unités de mesure et conversions", 2)
    add_p(doc, "Setup → Units of Measure + Unit Conversions : gramme, ml, chacun, caisse ; 1 lb = 453,592 g ; 1 oz liq = 29,5735 ml.")

    add_title(doc, "1.8 Préférences (Preferences)", 2)
    add_title(doc, "1.8.1 Inventory", 3)
    add_p(doc, INVENTORY_PREFS_NOTE)
    add_table(doc, INVENTORY_PREFS_HEADERS, inventory_prefs_rows())
    add_p(doc, "Sur les fiches Item : Key Item ON par défaut ; décochez sel/serviettes si besoin. Actualize ON uniquement sur Huile friture (item) — voir Partie 2.1 ; pas d'option Actualize sur l'écran Preps en v5.")
    add_title(doc, "1.8.2 Purchasing, Sales, Accounting", 3)
    add_bullets(doc, [
        "Purchasing : Show advanced tax adjustment fields → ON (obligatoire).",
        "Sales : dossier import POS si visible.",
        "Accounting : système comptable + dossier export (#49).",
    ])

    add_title(doc, "1.9 Import POS (Preferences → POS) et employés", 2)
    add_p(doc, POS_PREFS_NOTE)
    add_table(doc, POS_PREFS_HEADERS, pos_prefs_rows())
    add_p(doc, "Écran Import Specification Fields (engrenage, après Next) — fichier ventes_semaine_comptoir.csv :")
    add_table(doc, POS_SPEC_HEADERS, pos_spec_rows())
    add_bullets(doc, [
        "Security → Employees : Manager + Employé démo ; Access Levels configurés.",
    ])

    add_checklist(doc, [
        "Store / Company Information complété.",
        "Accounts (Name + Type + GL) + Categories + Inventory/Sales Groups.",
        "7 Storage Locations + 4 Suppliers + UOM.",
        "Taxes TCA + Tax Group TCA-ACHAT + Preferences (Purchasing advanced tax ON).",
        "Preferences POS (dossier CSV + spec) + 2 employés.",
        "CHECKPOINT : aucun item créé avant la Partie 2.",
    ])
    doc.add_page_break()

    # ---- PARTIE 2 ----
    add_title(doc, "Partie 2 — Créer les Items d'inventaire")
    add_p(doc, ITEM_CREATE_NOTE)
    add_table(doc, ITEM_CREATE_HEADERS, item_create_rows())
    add_p(doc, "Ordre de saisie : Core (gauche) → Case Size (droite) → Save → Locations / Conversions si besoin.")
    add_p(doc, CASE_SIZE_PURCHASE_UNIT_NOTE)
    add_p(doc, f"Sur chaque Case Size : Tax Group = {TAX_GROUP_CODE}, Price Includes Tax = décoché.")
    add_p(doc, "Sur la fiche Item (v5) : seulement Key Inventory Item et Actualize Usage Values — pas de Track Inventory ici.")
    add_p(doc, ITEM_MULTI_CASESIZE_NOTE)
    add_title(doc, "2.0 Trois premiers items (pain, bœuf, pommes)", 2)
    add_p(doc, ITEMS_FIRST3_NOTE)
    add_table(doc, ITEMS_FIRST3_CORE_HEADERS, items_first3_core_rows())
    add_table(doc, ITEMS_FIRST3_CASE_HEADERS, items_first3_case_rows())
    add_title(doc, "2.1 Items alimentaires — 13 suivants", 2)
    add_p(doc, ITEMS_REST13_NOTE)
    add_p(doc, ACTUALIZE_HUILE_NOTE)
    add_table(doc, ITEMS_DETAIL_CORE_HEADERS, items_rest13_core_rows())
    add_table(doc, ITEMS_DETAIL_CASE_HEADERS, items_rest13_case_rows())
    add_title(doc, "2.1b Exercice labo — Bacon, 2e fournisseur (Add case size)", 2)
    add_p(doc, ITEM_MULTI_CASESIZE_EXAMPLE_NOTE)
    add_table(doc, ITEM_MULTI_CASESIZE_EXAMPLE_HEADERS, items_multicasize_example_rows())
    add_title(doc, "2.2 Items boissons et papier", 2)
    add_p(doc, ITEMS_LAST5_NOTE)
    add_table(doc, ITEMS_DETAIL_CORE_HEADERS, items_last5_core_rows())
    add_table(doc, ITEMS_DETAIL_CASE_HEADERS, items_last5_case_rows())
    add_title(doc, "2.3 Questions de réflexion (Partie 2)", 2)
    add_bullets(doc, [
        "Pourquoi la laitue a un Yield de 75 % ?",
        "Quel est le coût au gramme du bœuf haché 80/20 ? (42 $US ÷ 5000 g)",
        "Pourquoi le bacon a 2 case sizes mais une seule fiche Core ?",
        "Pourquoi Huile friture a Actualize = Oui et le bœuf = Non ?",
        "À quoi sert la liste « Prix pour » (case, bag, each…) juste après Purchase Price ?",
        "Cochez Key Item sur le sel et les serviettes si vous voulez alléger le comptage.",
    ])
    add_checklist(doc, [
        "21 items créés sans erreur d'unité.",
        "Chaque item a un Primary Location.",
        "Bacon tranché : 2 case sizes (Distrib. Caraïbes + Boulangerie Pétion), View All = 2 lignes.",
        "Vous avez vérifié le coût unitaire calculé par OC sur 3 items au hasard.",
    ])
    doc.add_page_break()

    # ---- PARTIE 3 ----
    add_title(doc, "Partie 3 — Créer les Preps (préparations)")
    add_p(doc, "Recipe → Preps → New. Un Prep transforme des Items (ou d'autres Preps) en préparation réutilisable.")
    add_p(doc, PREPS_SCREEN_INTRO)
    add_title(doc, "3.0 Champs de l'écran New Prep", 2)
    add_p(doc, "Inventory Group pour tous les preps : Prep (défaut OC). Batch Uom : Batch.")
    add_table(doc, PREPS_FIELD_GUIDE_HEADERS, prep_field_guide_rows())
    add_p(doc, PREPS_NOTE)
    add_p(doc, PREPS_NESTED_NOTE)
    add_p(doc, PREPS_YIELD_NOTES)
    add_p(doc, PREPS_BATCH_UNIT_NOTE)
    add_p(doc, PREPS_ACTUALIZE_NOTE)
    add_title(doc, "3.0b Valeurs à saisir — Le Comptoir du Marché", 2)
    add_table(doc, PREPS_CORE_HEADERS, prep_core_rows())
    add_table(doc, PREPS_SUMMARY_HEADERS, prep_summary_rows())
    add_table(doc, PREPS_INGREDIENT_HEADERS, prep_ingredient_rows())
    add_table(doc, ["Prep", "Coût unitaire OC", "Calcul"], prep_cost_hint_rows())
    add_title(doc, "3.1 Boulettes burger (80 g crues)", 2)
    add_p(doc, "Voir tableau — 77 g bœuf haché 80/20 + 3 g sel fin → 1 boulette (80 g crue). Pas seulement le bœuf : assaisonnement obligatoire.")
    add_title(doc, "3.2 Frites maison — batch", 2)
    add_p(doc, "Voir tableau ingrédients — Pommes 5000 g + Huile 200 ml + Sel 50 g. Core : Batch Uom = Batch · Batch Yield = 1 batch · Recipe Unit = gram · Qty = 4000. Prep sheet ON, station Friterie.")
    add_title(doc, "3.2b Portion frites 200 g (prep dans prep)", 2)
    add_p(doc, "Créer APRÈS le batch. How is it used? = Unit. Batch Yield 1 each (= 200 g servies). Ingrédient : Add Prep → Frites maison — batch, 200 gram.")
    add_title(doc, "3.3 Sauce burger maison", 2)
    add_p(doc, "Mayonnaise 800 ml + Ketchup 200 ml. Core : Batch Uom = Batch · Batch Yield = 1 batch · Recipe Unit = ml · Qty = 950.")
    add_title(doc, "3.4 Chili — batch", 2)
    add_p(doc, "5 ingrédients — voir tableau. Core : Batch Uom = Batch · Batch Yield = 1 batch · Recipe Unit = gram · Qty = 12 000. Include prep sheet : Marmite.")
    add_checklist(doc, [
        "5 preps créés : tous les champs Core + Inventory + Prep Sheet selon tableaux 3.0.",
        "Portion frites utilise Add Prep → batch frites (prep dans prep).",
        "Actual Cost Per affiché et plausible sur chaque prep.",
        "Batch Yield et Recipe Unit cohérents (Qty recette calculée par OC).",
    ])
    doc.add_page_break()

    # ---- PARTIE 4 ----
    add_title(doc, "Partie 4 — Créer les Products (menu)")
    add_p(doc, "Recipe → Products → New. Un Product = plat vendu ; POS ID# lie les ventes importées.")
    add_title(doc, "4.0 POS ID# — liaison caisse ↔ Optimum Control", 2)
    add_p(doc, POS_ID_NOTE)
    add_p(doc, POS_ID_CHAIN_NOTE)
    add_p(doc, POS_ID_NOT_NOTE)
    add_p(doc, POS_ID_TILL_TAPE_NOTE)
    add_p(doc, POS_ID_FIELD_HINT)
    add_title(doc, "4.0a Même produit vendu sur plusieurs POS ?", 3)
    add_table(doc, POS_ID_MULTI_HEADERS, pos_id_multi_rows())
    add_title(doc, "4.0b Table PLU — exercice Le Comptoir", 3)
    add_table(doc, POS_ID_COMPTOIR_HEADERS, pos_id_comptoir_rows())
    add_p(doc, PRODUCTS_SCREEN_INTRO)
    add_title(doc, "4.0c Champs de l'écran New Product", 2)
    add_table(doc, PRODUCTS_FIELD_GUIDE_HEADERS, product_field_guide_rows())
    add_p(doc, PRODUCTS_NOTE)
    add_p(doc, PRODUCTS_NESTED_NOTE)
    add_p(doc, PRODUCTS_TOMATES_NOTE)
    add_title(doc, "4.0d Valeurs à saisir — Le Comptoir du Marché", 2)
    add_table(doc, PRODUCTS_CORE_HEADERS, product_core_rows())
    add_table(doc, PRODUCTS_INGREDIENT_HEADERS, product_ingredient_rows())
    add_title(doc, "4.1 Résumé menu", 2)
    add_table(doc,
        ["Product", "POS ID", "Prix vente", "Sales Category", "Ingrédients principaux"],
        [
            ["Burger classique", "101", "14,95 $US", "Food", "1 pain, 1 boulette, 15 ml sauce, 30 g tomates, 1 boîte, 1 serviette"],
            ["Burger bacon", "102", "16,95 $US", "Food", "Burger classique + 2 tranches bacon"],
            ["Frites moyennes", "201", "5,50 $US", "Food", "1 Portion frites 200 g (prep), 1 serviette"],
            ["Salade César", "301", "13,50 $US", "Food", "150 g laitue, 45 ml dressing, parmesan, croûtons"],
            ["Bol chili", "401", "9,95 $US", "Food", "400 g chili, 1 fourchette"],
            ["Cola", "501", "3,25 $US", "Beverage", "1 canette Cola 355 ml"],
            ["Eau", "502", "2,75 $US", "Beverage", "1 bouteille Eau 500 ml"],
        ])
    add_title(doc, "4.2 Burger classique (exemple pas à pas)", 2)
    add_p(doc, "Sales Group = Entrées / Plats (obligatoire). POS ID# 101 · Selling Price 14,95 $US. Ingrédients : pain, boulette, sauce, Tomates 30 gram (pas en tranches), boîte, serviette.")
    add_title(doc, "4.3 Burger bacon — product dans product", 2)
    add_p(doc, "Add Product → Burger classique (1 each) + Add Item → Bacon tranché (2 each).")
    add_title(doc, "4.4 Frites moyennes (POS 201)", 2)
    add_table(doc, ["Champ", "Valeur"], [
        ["Description", "Frites moyennes"],
        ["Sales Group", "Accompagnements"],
        ["POS ID#", "201"],
        ["Selling Price", "5,50 $US"],
        ["Tax Group", "No Tax Group"],
    ])
    add_p(doc, "Ordre dans OC : Recipe → Products → New → remplir Core + Pricing → ajouter 2 lignes ingrédients (Add Prep puis Add Item) :")
    add_table(doc, ["Étape", "Bouton", "Ingrédient", "Qty", "Unité"], [
        ["1", "Add Prep", "Portion frites 200 g", "1", "each"],
        ["2", "Add Item", "Serviette", "1", "each"],
    ])
    add_note(doc, "Utiliser le prep Portion frites 200 g — pas Frites maison — batch. Save → Cost Percent ≈ 5–6 %.")

    add_title(doc, "4.5 Salade César (POS 301)", 2)
    add_p(doc, SALADE_CESAR_NOTE)
    add_table(doc, ["Champ", "Valeur"], [
        ["Description", "Salade César"],
        ["Sales Group", "Entrées / Plats"],
        ["POS ID#", "301"],
        ["Selling Price", "13,50 $US"],
        ["Tax Group", "No Tax Group"],
    ])
    add_p(doc, "Tous les ingrédients sont des Items (pas de prep). Add Item pour chaque ligne — **4 items**, unités en gram ou ml :")
    add_table(doc, ["Ordre", "Ingrédient (Item)", "Qty", "Unité", "Rappel item Partie 2"],
        [
            ["1", "Laitue romaine", "150", "gram", "Produce · yield 75 %"],
            ["2", "César dressing", "45", "ml", "Dairy · Volume"],
            ["3", "Parmesan râpé", "15", "gram", "Dairy · Weight"],
            ["4", "Croûtons", "20", "gram", "Dry Goods · Weight"],
        ])
    add_p(doc, "Vérification : Actual Product Cost se remplit après les 4 lignes (~1,43 $US). Aucune serviette/boîte sur cette salade dans l'exercice.")

    add_title(doc, "4.6 Bol chili (POS 401)", 2)
    add_table(doc, ["Champ", "Valeur"], [
        ["Description", "Bol chili"],
        ["Sales Group", "Entrées / Plats"],
        ["POS ID#", "401"],
        ["Selling Price", "9,95 $US"],
        ["Tax Group", "No Tax Group"],
    ])
    add_table(doc, ["Étape", "Bouton", "Ingrédient", "Qty", "Unité"], [
        ["1", "Add Prep", "Chili — batch", "400", "gram"],
        ["2", "Add Item", "Fourchette plastique", "1", "each"],
    ])
    add_note(doc, "400 gram = portion bol (prep batch yield total = 12 000 g). Prep Chili — batch doit exister (Partie 3).")

    add_title(doc, "4.7 Cola et Eau (POS 501–502)", 2)
    add_p(doc, "Boissons = 1 item direct par product, Sales Group Boissons. Pas de prep.")
    add_table(doc, ["Product", "POS ID#", "Selling Price", "Add Item", "Qty", "Unité"], [
        ["Cola", "501", "3,25 $US", "Cola 355 ml", "1", "each"],
        ["Eau", "502", "2,75 $US", "Eau 500 ml", "1", "each"],
    ])

    add_title(doc, "4.8 Vérifier les coûts recette", 2)
    add_p(doc, "Pour chaque Product, notez le food cost % affiché par OC :")
    add_table(doc, ["Product", "Coût recette", "Prix vente", "Food cost %", "Cible 32 % ?"], [
        ["Burger classique", "à compléter", "14,95 $US", "", ""],
        ["Burger bacon", "à compléter", "16,95 $US", "", ""],
        ["Frites moyennes", "à compléter", "5,50 $US", "", ""],
        ["Salade César", "à compléter", "13,50 $US", "", ""],
        ["Bol chili", "à compléter", "9,95 $US", "", ""],
    ])
    add_checklist(doc, [
        "7 products créés avec POS ID unique.",
        "Tous les products sont liés (état prêt pour import POS).",
        "Food cost % noté pour chaque plat.",
    ])
    doc.add_page_break()

    # ---- PARTIE 5 ----
    add_title(doc, f"Partie 5 — Inventaire d'ouverture (Lundi {EXO_DATE_OPENING})")
    add_p(doc, "(Parcours Linéaire : étape 27 — Phase D.)")
    add_p(doc, OPENING_INVENTORY_NOTE)
    add_p(doc, COUNTSHEET_TRACK_NOTE)
    add_p(doc, "Ordre dans cet exercice (important) :")
    add_bullets(doc, OPENING_INVENTORY_FLOW)
    add_p(doc, "Écran Create Inventory (Inventory → Count Inventory → New) — vidéo #25 :")
    add_table(doc, OPENING_INVENTORY_WIZARD_HEADERS, opening_inventory_wizard_rows())
    add_p(doc, OPENING_INVENTORY_COUNT_NOTE)
    add_p(doc, "Après Finish : tri Location → saisir les quantités du tableau (Purchase Count en cs ou batch ; Pak Count si indiqué). Aucune facture requise :")
    add_table(doc, OPENING_INVENTORY_QTY_HEADERS, opening_inventory_qty_rows())
    add_step(doc, 1, "Tri recommandé : Location (ou Custom si étape 26 faite).")
    add_step(doc, 2, "Saisissez toutes les lignes → Summarize Count → Finalize.")
    add_step(doc, 3, f"Notez la valeur totale d'inventaire d'ouverture : _________________ {DEVISE_COURTE}")
    add_title(doc, "5.1 Où voir l'inventaire et Qty on Hand (OC v5)", 2)
    add_p(doc, OPENING_INVENTORY_VIEW_NOTE)
    add_table(doc, OPENING_INVENTORY_VIEW_HEADERS, opening_inventory_view_rows())
    add_checklist(doc, [
        f"Inventaire d'ouverture Finalize (date = {EXO_DATE_OPENING}, lundi soir).",
        "Items : Qty on Hand vérifiée sur Inventory → Items (ex. Pain burger = 24 each).",
        "Preps : pas de Qty on Hand sur fiche Prep — feuille Count Inventory ou Reports → Prep Sheet (On Hand).",
        "Aucun item négatif après Summarize.",
        "Aucune facture saisie avant cette étape — normal.",
    ])
    doc.add_page_break()

    # ---- PARTIE 6 ----
    add_title(doc, "Partie 6 — Semaine d'exploitation (Mardi → Samedi)")
    add_p(doc, "(Parcours Linéaire : étapes 29–36 — Phase E puis Phase F.)")
    add_p(doc, "Prérequis : inventaire d'ouverture Finalize (Partie 5). Les factures augmentent le stock déjà posé à l'ouverture — elles ne le remplacent pas.")
    add_p(doc, "Respectez l'ordre chronologique : Factures → Ventes (Sales Mix) → Daily Sales → Waste.")

    add_title(doc, "6.1 Factures manuelles (Purchasing → Invoices)", 2)
    add_p(doc, EXO_WEEK_NOTE)
    add_table(doc, EXO_WEEK_HEADERS, exo_week_rows())
    add_p(doc, INVOICE_PHASE_NOTE)
    add_p(doc, INVOICE_MANUAL_NOTE)
    add_p(doc, "Ordre de saisie dans OC :")
    add_bullets(doc, INVOICE_ENTRY_FLOW)
    add_table(doc, INVOICE_SCREEN_HEADERS, invoice_screen_rows())
    add_title(doc, "6.1.1 Facture labo — Boulangerie Pétion (mardi 30/06)", 3)
    add_p(doc, INVOICE_LAB_BOULANGERIE)
    add_title(doc, "6.1.2 Facture mardi — Distrib. Caraïbes", 3)
    add_p(doc, INVOICE_LAB_DISTRIB_LUNDI)
    add_title(doc, "6.1.3 Toutes les lignes de la semaine", 3)
    add_table(doc, INVOICE_LINE_HEADERS, invoice_line_rows())
    add_table(doc, INVOICE_SUMMARY_HEADERS, invoice_summary_rows())
    add_p(doc, "Invoice Total dans OC = **Total TTC** de la colonne de droite. Account Balance = 0 avant Save.")
    add_checklist(doc, [
        "5 factures enregistrées (BP, DC×2, BN, EH).",
        "TCA 10 % sur chaque facture (onglet Taxes and Adjustments).",
        "Qty on Hand vérifiée sur 2 items après la 1re facture.",
    ])

    add_title(doc, "6.2 Commande (option Phase E — étape 30)", 2)
    add_table(doc, ORDER_MARDI_HEADERS, order_mardi_rows())
    add_p(doc, "Purchasing → Orders → New. Même logique que facture : Supplier, items, Save.")

    add_title(doc, "6.3 Ventes POS — Till Tape (mar.) puis import CSV (mer→sam)", 2)
    add_p(doc, TILL_TAPE_NOTE)
    add_p(doc, TILL_TAPE_TWO_STEP_NOTE)
    add_title(doc, f"6.3.1 Till Tape — mardi {EXO_DAY_TO_DATE[EXO_FIRST_SALES_DAY]} (Phase F, étape 32 · vidéo #43)", 3)
    add_p(doc, "Ordre obligatoire : d'abord la **liste**, ensuite la **vente du jour**.")
    add_bullets(doc, TILL_TAPE_LIST_FLOW)
    add_bullets(doc, TILL_TAPE_SALE_FLOW)
    add_table(doc, TILL_TAPE_QTY_HEADERS, till_tape_mardi_rows())
    add_p(doc, TILL_TAPE_POS_NOTE)
    add_checklist(doc, TILL_TAPE_VALIDATE)

    add_title(doc, "6.3.2 Import CSV — mardi à vendredi (étape 33 · vidéo #42)", 3)
    add_bullets(doc, [
        "Garder **Comptoir CSV** dans Settings → Preferences → POS.",
        "Sales → New → Import From File → Browse → ventes_mercredi_2026-07-01.csv … ventes_samedi_2026-07-04.csv.",
        "Fichiers **5 colonnes** (Date en col. 1) — même format que ventes_semaine_comptoir.csv.",
        "Ne pas importer ventes_minimal_*.csv avec Comptoir CSV (2 colonnes → erreur d'index).",
        "Alternative : ventes_semaine_comptoir.csv (semaine complète) ou Till Tape chaque jour.",
        "Traitez Pending Sales (Unlinked / Mismatched).",
    ])
    add_p(doc, CSV_IMPORT_NOTE)
    add_table(doc, CSV_IMPORT_HEADERS, csv_import_map_rows())
    add_step(doc, 1, "Vérifiez Preferences → POS : dossier ventes_csv + Comptoir CSV (engrenage : indices 1→5 comme tableau ci-dessus).")
    add_p(doc, "Quantités à importer (si saisie manuelle) :")
    add_table(doc, ["Date", "Jour", "101 Burger", "102 Bacon", "201 Frites", "301 César", "401 Chili", "501 Cola", "502 Eau"], [
        ["30/06/2026", "Mardi", "25", "10", "30", "8", "15", "40", "12"],
        ["01/07/2026", "Mercredi", "28", "12", "35", "10", "18", "45", "15"],
        ["02/07/2026", "Jeudi", "22", "8", "28", "12", "20", "38", "10"],
        ["03/07/2026", "Vendredi", "30", "15", "40", "9", "16", "50", "18"],
        ["04/07/2026", "Samedi", "35", "18", "45", "14", "22", "55", "20"],
    ])
    add_step(doc, 1, "Après import, ouvrez Pending Sales (étape 6.3.3) si des lignes sont Unlinked.")
    add_step(doc, 2, "Vérifiez que chaque POS ID est Linked ou Valid.")

    add_title(doc, "6.3.3 Pending Sales — lier les ventes POS (étape 34 · vidéo #44)", 3)
    add_p(doc, PENDING_SALES_NOTE)
    add_p(doc, POS_ID_CHAIN_NOTE)
    add_p(doc, PENDING_SALES_UNLINK_NOTE)
    add_p(doc, PENDING_SALES_COMPTOIR_NOTE)
    add_bullets(doc, PENDING_SALES_FLOW)
    add_table(doc, PENDING_SALES_STATES_HEADERS, pending_sales_state_rows())
    add_checklist(doc, PENDING_SALES_VALIDATE)

    add_title(doc, "6.4 Ventes quotidiennes (Daily Sales)", 2)
    add_p(doc, DAILY_SALES_NOTE)
    add_p(doc, DAILY_SALES_SCREEN_NOTE)
    add_p(doc, DAILY_SALES_GROSS_NOTE)
    add_p(doc, DAILY_SALES_DISTRIB_NOTE)
    add_bullets(doc, DAILY_SALES_FLOW)
    add_title(doc, "6.4.1 Lexique — champs de l'écran Daily Sales", 3)
    add_table(doc, DAILY_SALES_FIELDS_HEADERS, daily_sales_field_rows())
    add_title(doc, "6.4.2 Valeurs par jour (réparties sur la semaine)", 3)
    add_table(doc, DAILY_SALES_OPS_HEADERS, daily_sales_ops_rows())
    add_title(doc, "6.4.3 Résumé rapide", 3)
    add_table(doc, DAILY_SALES_TABLE_HEADERS, daily_sales_rows())
    add_checklist(doc, DAILY_SALES_VALIDATE)

    add_title(doc, "6.5 Pertes (Waste)", 2)
    add_p(doc, WASTE_NOTE)
    add_bullets(doc, WASTE_FLOW)
    add_table(doc, WASTE_TABLE_HEADERS, waste_rows())
    add_checklist(doc, WASTE_VALIDATE)
    add_checklist(doc, [
        "5 factures saisies et équilibrées.",
        "5 jours de Sales Mix saisis.",
        "5 Daily Sales saisies.",
        "4 entrées Waste saisies.",
    ])
    doc.add_page_break()

    # ---- PARTIE 7 ----
    add_title(doc, f"Partie 7 — Inventaire de clôture (Samedi {EXO_DATE_CLOSING})")
    add_p(doc, CLOSING_INVENTORY_NOTE)
    add_bullets(doc, CLOSING_INVENTORY_FLOW)
    add_table(doc, CLOSING_INVENTORY_QTY_HEADERS, closing_inventory_qty_rows())
    add_p(doc, "Complétez le reste des items avec des quantités cohérentes avec vos achats et ventes.")
    add_checklist(doc, CLOSING_INVENTORY_VALIDATE)
    doc.add_page_break()

    # ---- PARTIE 8 ----
    add_title(doc, "Partie 8 — Commandes et Par Levels")
    add_title(doc, "8.1 Item Par Levels", 2)
    add_p(doc, "Inventory → Item Par Levels. Configurez :")
    add_table(doc, ["Item", "Min (reorder)", "Max"], [
        ["Pain burger", "12", "36"],
        ["Bœuf haché 80/20", "5000 g", "15000 g"],
        ["Cola 355 ml", "24", "72"],
    ])
    add_step(doc, 1, "Vérifiez la tuile Order Reminder sur le Dashboard : des items doivent apparaître sous le minimum après votre semaine.")
    add_step(doc, 2, "Créez une commande depuis Order Reminder → Round to Case → Save (sans exporter si exercice seul).")

    add_title(doc, "8.2 Prep Par Levels (optionnel)", 2)
    add_p(doc, "Sur le prep Chili : Include on prep sheets, shelf life 3 jours. Recipe → Prep Par Levels → Calculate Required Amounts sur la semaine. Imprimez Prep Sheet Daily pour un jour au choix.")
    doc.add_page_break()

    # ---- PARTIE 9 ----
    add_title(doc, "Partie 9 — Rapports et analyse")
    add_title(doc, "9.1 Usage Summary (rapport central)", 2)
    add_p(doc, USAGE_SUMMARY_NOTE)
    add_p(doc, USAGE_SUMMARY_FORMULA)
    add_bullets(doc, USAGE_SUMMARY_FLOW)
    add_step(doc, 1, "Répondez par écrit :")
    add_bullets(doc, USAGE_SUMMARY_QUESTIONS)
    add_checklist(doc, USAGE_SUMMARY_VALIDATE)

    add_title(doc, "9.2 Autres rapports et backup", 2)
    add_p(doc, BACKUP_NOTE)
    add_bullets(doc, BACKUP_FLOW)
    add_checklist(doc, [
        "Cost of Sales Analysis généré.",
        "Item Activity Report sur le bœuf haché.",
        "Recipe Book PDF exporté (au moins Burgers + Chili).",
        "Export Excel depuis Usage Summary testé.",
    ])
    doc.add_page_break()

    # ---- PARTIE 11 ----
    add_title(doc, "Partie 11 — Multi-stocks : emplacements et transferts")
    add_p(doc, "Cette partie couvre deux niveaux : (A) plusieurs emplacements dans un même restaurant, et (B) transferts entre deux établissements (OC Premier / multi-magasins).")

    add_title(doc, "11.A Réorganiser le stock entre emplacements (un seul magasin)", 2)
    add_p(doc, f"Situation : la livraison {NOM_DISTRIB} arrive en Réserve ; vous devez approvisionner la Ligne chaude et le Froid avant le service.")
    add_table(doc, ["Action terrain", "Ce que vous faites dans OC", "Remarque"], [
        ["Déplacer 10 kg bœuf Congélateur → Ligne chaude (décongélation)", "Pas de transfert formel : le stock reste sous le même item. Comptez au bon endroit à l'inventaire.", "OC suit la quantité totale ; les emplacements guident le comptage."],
        ["Remplir le petit bac ketchup (Sec → Ligne chaude)", "Secondary Location sur l'item ; à l'inventaire, répartissez les quantités par emplacement si votre feuille le permet.", "Le Primary reste « Cuisine — Sec » pour les commandes."],
        ["Corriger un déplacements oublié avant inventaire", "Inventory → Adjust Inventory sur l'item concerné.", "Ajustez la quantité globale si le déplacement a été consommé ou perdu."],
        ["Compter shelf-to-sheet", "Count Inventory → Customize Sort par emplacement.", "Réfrigérateur de haut en bas = ordre Custom."],
    ])
    add_step(doc, 1, "Mardi matin (simulation) : notez sur papier que vous « déplacez » 5 kg de bœuf de Congélateur vers Ligne chaude.")
    add_step(doc, 2, "À l'inventaire de clôture, comptez le bœuf dans le Congélateur et à la Ligne chaude séparément (tri Location).")
    add_step(doc, 3, "Vérifiez que la somme des emplacements = quantité totale de l'item.")
    add_checklist(doc, [
        "Inventaire de clôture saisi avec tri Location.",
        "Customize Sort configuré pour au moins un emplacement.",
        "Vous comprenez la différence Primary (commandes) vs Secondary (organisation).",
    ])

    add_title(doc, "11.B Transferts entre établissements (OC Premier — optionnel)", 2)
    add_note(doc, "Nécessite une base multi-magasins avec au moins 2 établissements. Si vous n'avez qu'OC Pro mono-magasin, lisez cette section à titre informatif.")
    add_p(doc, "Scénario étendu : ajoutez un second établissement « Entrepôt Le Comptoir » (stock central) en plus du restaurant « Le Comptoir du Marché ».")
    add_table(doc, ["Établissement", "Rôle", "Stock initial"], [
        ["Entrepôt Le Comptoir", f"Réserve centrale — grosses livraisons {NOM_DISTRIB}", "Bœuf 20 kg, pommes 40 kg, cola 96 canettes"],
        ["Le Comptoir du Marché", "Point de vente — cuisine active", "Stock cuisine minimal"],
    ])
    add_p(doc, "Partagez les items entre magasins : Settings → Utilities → Item and Store Management → Send Items vers le restaurant.")
    add_title(doc, "Méthode 1 — Réquisition (demande de stock)", 3)
    add_step(doc, 1, "Connecté au RESTAURANT : Stock Transfer → Request Stock → New.")
    add_step(doc, 2, "Sélectionnez « Entrepôt Le Comptoir » comme source.")
    add_step(doc, 3, "Demandez : Bœuf 3000 g, Pommes de terre 10000 g, Cola 24 canettes → Save.")
    add_step(doc, 4, "Basculez sur l'ENTREPÔT : Stock Transfer → Approve Requests → ouvrez la demande → Approve.")
    add_step(doc, 5, "Fulfill Requisitions → cochez la demande → vérifiez On Hand → Save.")
    add_note(doc, "Si le stock est insuffisant, OC propose un backorder pour le solde non livré (ex. : demandé 5 kg, disponible 3,25 kg).")
    add_title(doc, "Méthode 2 — Transfert direct (sans réquisition)", 3)
    add_step(doc, 1, "Depuis l'ENTREPÔT (magasin qui envoie) : Stock Transfer → Transfers → New.")
    add_step(doc, 2, "Destination : Le Comptoir du Marché.")
    add_step(doc, 3, "Items : Bacon 500 g, Huile friture 2000 ml → Save and Extract.")
    add_checklist(doc, [
        "Réquisition créée, approuvée et fulfillée (ou transfert direct réussi).",
        "Stock diminué à l'entrepôt et augmenté au restaurant.",
        "Usage Summary consolidé (Enterprise) ou par magasin selon votre config.",
    ])
    doc.add_page_break()

    # ---- PARTIE 10 ----
    add_title(doc, "Partie 10 — Défis bonus (dépannage)")
    add_p(doc, "Réalisez ces scénarios pour tester votre compréhension :")
    add_title(doc, "Défi A — Erreur de comptage", 2)
    add_p(doc, "Vous avez saisi 2 cs de pommes de terre au lieu de 1 cs à l'ouverture (≈45 kg au lieu de ≈23 kg). Corrigez via Inventory → Adjust Inventory sans refaire tout l'inventaire.")
    add_title(doc, "Défi B — Recette modifiée", 2)
    add_p(doc, "Ajoutez 10 g de bacon dans le Burger classique. Utilisez Rerun Sales Mix sur la semaine pour recalculer l'usage idéal.")
    add_title(doc, "Défi C — Item désactivé", 2)
    add_p(doc, "Désactivez l'Eau 500 ml (plus vendue). Vérifiez qu'elle n'apparaît plus aux commandes mais reste en historique.")
    add_title(doc, "Défi D — Réévaluation", 2)
    add_p(doc, "Le prix du bœuf passe de 42 $US à 48 $US la caisse sur la dernière facture. Appliquez la procédure Revaluing Inventory du chapitre 6 du guide.")
    doc.add_page_break()

    # ---- CORRIGÉ ----
    add_title(doc, "Annexe A — Résultats attendus (résumé)")
    add_p(doc, "Corrigé détaillé avec tous les calculs : voir Exercice_Corrige_Detaille.docx dans le même dossier.")
    add_p(doc, "Ordres de grandeur — vos chiffres peuvent varier légèrement selon arrondis et taxes.")
    add_title(doc, "Coûts recette approximatifs", 2)
    add_table(doc, ["Product", "Coût estimé", "Food cost %"], [
        ["Burger classique", "~4,80 $US", "~32 %"],
        ["Burger bacon", "~5,60 $US", "~33 %"],
        ["Frites moyennes", "~1,65 $US", "~30 %"],
        ["Salade César", "~1,43 $US", "~31 %"],
        ["Bol chili", "~3,10 $US", "~31 %"],
    ])
    add_title(doc, "Usage Summary — ce que vous devriez observer", 2)
    add_bullets(doc, [
        "Huile friture : écart possible si Actualize non activé ou frites en perte.",
        "Laitue : écart lié au waste mardi + yield 75 %.",
        "Chili : usage idéal suit les ventes du bol chili ; perte jeudi augmente l'actual.",
        "Pain : ventes burgers + waste vendredi = cohérence à vérifier.",
    ])
    add_title(doc, "Critères de réussite globale", 2)
    add_checklist(doc, [
        "Aucune vente Unlinked en fin d'exercice.",
        "Usage Summary généré sans message « Sales not defined » sur Food.",
        "Inventaire clôture − ouverture + achats ≈ usage actual (ordre de grandeur).",
        "Vous savez expliquer un écart à un propriétaire en 3 phrases.",
        "Sauvegarde Backup Data effectuée en fin de session.",
    ])

    add_title(doc, "Annexe B — Routine hebdomadaire (après l'exercice)")
    add_table(doc, ["Jour", "Tâches OC"], [
        ["Lundi", "Factures weekend, import ventes, Daily Sales, waste"],
        ["Mardi", "Commandes (Order Reminder), réception factures"],
        ["Mercredi", "Prep par levels, prep sheets"],
        ["Jeudi", "Factures, rapports intermédiaires"],
        ["Vendredi", "Import ventes, Daily Sales, backup"],
        ["Samedi", "Inventaire complet, Usage Summary, analyse écarts"],
        ["Dimanche", "Planification commandes, review food cost"],
    ])

    add_title(doc, "Annexe C — Correspondance avec le guide")
    add_table(doc, ["Partie exercice", "Chapitre du guide"], [
        ["Partie 1", "Ch. 1–2, 12"],
        ["Partie 1.5", "Ch. 3 — Storage Locations, Ch. 6 — tri Location"],
        ["Partie 2", "Ch. 3"],
        ["Partie 3–4", "Ch. 4"],
        ["Partie 5–7", "Ch. 6, 9, 10"],
        ["Partie 6 factures", "Ch. 8"],
        ["Partie 6.2 CSV", "Ch. 9 — Import POS, Pending Sales"],
        ["Partie 8", "Ch. 7, Prep Par Levels"],
        ["Partie 9", "Ch. 13"],
        ["Partie 11.A", "Ch. 3, 6 — emplacements et comptage"],
        ["Partie 11.B", "Ch. 2 — Stock Transfer, Item Store Management"],
        ["Défis", "Ch. 6, 9, 18, 20"],
    ])

    add_title(doc, "Annexe D — Fichiers fournis")
    add_table(doc, ["Fichier / dossier", "Description"], [
        ["exercice_comptoir/ventes_csv/", "11 fichiers CSV + référence Till Tape"],
        ["exercice_comptoir/README_import_ventes.md", "Mode d'emploi import ventes"],
        ["Exercice_Corrige_Detaille.docx", "Corrigé complet avec calculs"],
    ])

    publish_docx(doc, OUT)

if __name__ == "__main__":
    build()
