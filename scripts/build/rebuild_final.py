#!/usr/bin/env python3
"""Reconstruction correcte : nouveau contenu à la FIN de chaque section, pas sous le titre."""
import re, shutil, sys, zipfile
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

sys.path.insert(0, str(Path(__file__).resolve().parent))
from exercice_locale import (
    ITEM_CREATE_DOC, ITEM_CREATE_NOTE,
    item_multicasize_book_blocks,
    PREPS_CREATE_NOTE, PREPS_CREATE_DOC, prep_book_blocks,
    PRODUCTS_CREATE_NOTE, PRODUCTS_CREATE_DOC, product_book_blocks,
    pos_id_book_blocks, POS_ID_NOTE,
    opening_inventory_book_blocks,
    invoice_book_blocks,
    sales_till_tape_book_blocks,
    daily_sales_book_blocks,
    POS_SPEC_DOC, POS_SPEC_WIZARD_DOC,
    TAX_GROUP_CODE, TAX_GROUPS_NOTE, TAX_GROUP_DOC,
)

SRC = Path(r"E:\OC DOCS\Comprendre_Optimum_Control_backup.docx")
DOC = Path(r"E:\OC DOCS\Comprendre_Optimum_Control_V2.docx")
PLAT = Path(r"E:\OC DOCS\Comprendre_Optimum_Control_V2_PLAT.docx")

def hlevel(sn):
    if not sn or not sn.startswith("Heading"): return 99
    n = sn.replace("Heading ", "")
    return int(n) if n.isdigit() else 99

def pstyle(style, text):
    if style == "Normal":
        return "List Paragraph" if text.strip().startswith("•") else None
    return style

def insert_before(anchor, text, style):
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    s = pstyle(style, text)
    if s:
        try: para.style = s
        except KeyError: pass
    if text: para.add_run(text)
    return para

def insert_block_before(anchor, blocks):
    for st, tx in blocks:
        anchor = insert_before(anchor, tx, st)
    return anchor

def insert_after(anchor, text, style):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    s = pstyle(style, text)
    if s:
        try: para.style = s
        except KeyError: pass
    if text: para.add_run(text)
    return para

def insert_block_after(anchor, blocks):
    cur = anchor
    for st, tx in blocks:
        cur = insert_after(cur, tx, st)
    return cur

def find_p(doc, needle, after=None):
    ok = after is None
    for p in doc.paragraphs:
        if after and after in p.text: ok = True; continue
        if ok and needle in p.text: return p
    return None

def pidx(doc, target):
    for i, p in enumerate(doc.paragraphs):
        if p._p is target._p: return i
    return -1

def next_sibling_heading(doc, anchor):
    """Prochain titre de niveau égal ou supérieur (fin de section)."""
    idx = pidx(doc, anchor)
    lv = hlevel(anchor.style.name if anchor.style else "")
    for j in range(idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[j]
        sn = p.style.name if p.style else ""
        if sn.startswith("Heading") and hlevel(sn) <= lv:
            return p
    return doc.paragraphs[-1]

def add_at_end_of_section(doc, section_heading, blocks, after=None):
    """Ajoute blocks à la fin de la section, après son texte existant."""
    anchor = find_p(doc, section_heading, after)
    if not anchor:
        print("MANQUE", section_heading)
        return False
    end_heading = next_sibling_heading(doc, anchor)
    idx = pidx(doc, end_heading)
    if idx <= 0:
        print("ERREUR idx", section_heading)
        return False
    last = doc.paragraphs[idx - 1]
    while not last.text.strip() and idx > 1:
        idx -= 1
        last = doc.paragraphs[idx - 1]
    insert_block_after(last, blocks)
    print("OK fin section", section_heading[:48], "+", len(blocks))
    return True

def insert_right_after_heading(doc, section_heading, blocks, after=None):
    """Insère du texte juste sous le titre, avant tableau ou section suivante."""
    anchor = find_p(doc, section_heading, after)
    if not anchor:
        print("MANQUE intro", section_heading)
        return False
    insert_block_after(anchor, blocks)
    print("OK intro", section_heading[:48], "+", len(blocks))
    return True

def rows_to_blocks(heading2, intro, rows):
    """Convertit des lignes (zone, champ, rôle, note) en paragraphes Word."""
    blocks = [("Heading2", heading2)]
    if intro:
        blocks.append(("Normal", intro))
    zone = None
    for row in rows:
        z, champ, role = row[0], row[1], row[2]
        hint = row[3] if len(row) > 3 else ""
        if z != zone:
            zone = z
            blocks.append(("Heading3", zone))
        line = f"• {champ} — {role}"
        if hint:
            line += f" ({hint})"
        blocks.append(("Normal", line))
    return blocks

ITEM_CREATE_BLOCKS = rows_to_blocks(
    "Référence complète — Items → New (vidéo #06)",
    ITEM_CREATE_NOTE + f" Tax Group exercice : {TAX_GROUP_CODE} sur chaque Case Size.",
    ITEM_CREATE_DOC,
)

PREPS_CREATE_BLOCKS = rows_to_blocks(
    "Référence complète — Preps → New (vidéo #16)",
    PREPS_CREATE_NOTE,
    PREPS_CREATE_DOC,
)

PRODUCTS_CREATE_BLOCKS = rows_to_blocks(
    "Référence complète — Products → New (vidéo #20–21)",
    PRODUCTS_CREATE_NOTE,
    PRODUCTS_CREATE_DOC,
)

TAX_GROUP_EXERCICE_BLOCKS = rows_to_blocks(
    f"Tax Group — exercice Comptoir ({TAX_GROUP_CODE})",
    TAX_GROUPS_NOTE,
    [("Groupe", c, v, u) for c, v, u, _ in [(row[0], row[1], row[2], row[3]) for row in [
        (TAX_GROUP_CODE, "Achats soumis à TCA", "TCA 10 %", "Assigner à tous les items ; Price Includes Tax = Non"),
    ]]],
)

POS_SPEC_WIZARD_BLOCKS = rows_to_blocks(
    "Assistant Import Specification — écran 1 (engrenage)",
    "Settings → Preferences → POS → engrenage à côté de POS configuration. "
    "Renseignez ces champs puis Next > pour l'écran Import Specification Fields.",
    [( "Écran 1", c, v, "") for c, v in POS_SPEC_WIZARD_DOC ],
)

POS_SPEC_FIELD_BLOCKS = rows_to_blocks(
    "Import Specification Fields — écran 2 (Index = n° de colonne)",
    "Fichier ventes_semaine_comptoir.csv : colonne 1=Date, 2=POS ID, 3=Description, "
    "4=Quantity, 5=Sales Price. Champs * obligatoires.",
    [("Écran 2", c, v, "") for c, v, _ in POS_SPEC_DOC],
)

def add_before_conclusion(doc, blocks):
    concl = find_p(doc, "Conclusion")
    idx = pidx(doc, concl)
    insert_block_after(doc.paragraphs[idx - 1], blocks)
    print("OK chapitres 19-20 +", len(blocks))

# -------- CONTENU ENRICHI --------
DASHBOARD = [
    ("Heading2", "Le tableau de bord (Dashboard)"),
    ("Normal", "Depuis la version 5, Optimum Control ouvre sur un tableau de bord personnalisable. Ajoutez des tuiles via l'icône + verte."),
    ("Normal", "Tuiles utiles : Order Reminder (items sous le par minimum, bouton Create Order), Top 10 Price Changes, Period Purchases by Supplier or Group, General Activity."),
    ("Heading3", "Autres nouveautés de l'interface (version 5)"),
    ("Normal", "Listes calendrier dans Invoices, Sales et Waste ; rapports favoris par utilisateur ; export Excel direct depuis les rapports ; drill-downs Usage Summary enrichis (fiche item, comptages, factures, usage idéal)."),
]

DEACTIVATE = [
    ("Heading2", "Désactiver un item, prep ou product"),
    ("Normal", "Quand un élément n'est plus utilisé : Miscellaneous → Item and Recipe Management → onglet concerné → décochez Active → Save."),
    ("Normal", "Si l'item est encore dans une recette (ligne en gras), retirez-le des recettes ou désactivez les recettes parentes d'abord."),
    ("Normal", "Pour réactiver : filtre → Inactive items → recochez Active (élément en rouge). Après Amalgamate : désactivez le doublon fusionné."),
]

BARCODE = [
    ("Heading3", "Code-barres et OC Mobile"),
    ("Normal", "Le champ Barcode sur Item et Prep sert au scan avec OC Mobile. Le champ OR Code sert à l'import EDI fournisseur."),
]

PREP_PAR = [
    ("Heading2", "Niveaux Par des preps (Prep Par Levels)"),
    ("Normal", "Calcule combien de préparation produire selon les ventes importées et la durée de conservation (shelf life) de chaque prep."),
    ("Heading3", "Configuration sur la fiche Prep"),
    ("Normal", "Section Prep Sheet Information : cochez Include on prep sheets, indiquez les stations et la shelf life (ex. 6 jours)."),
    ("Heading3", "Fenêtre Prep Par Levels"),
    ("Normal", "Recipe → Prep Par Levels → incluez les preps → Calculate Required Amounts → choisissez la plage de dates → ajustez min/max par jour → Save."),
    ("Heading3", "Rapports prep sheet"),
    ("Normal", "Reports → Prep Sheet Daily ou Weekly. Colonnes On Hand (stock actuel) et Make (à produire). Prep Margin % = tampon (défaut 10 %)."),
]

RECIPE_BOOK = [
    ("Heading2", "Livre de recettes (Recipe Book)"),
    ("Normal", "Reports → Recipe Book → onglets Preps et Products → sélectionnez → View Book → imprimez ou enregistrez en PDF."),
]

MENU_REPORT = [
    ("Heading2", "Rapport Menu Product Detail Problems"),
    ("Normal", "Reports → Recipe → Menu Product Detail Problems → seuil de coût (ex. 30 %) → Run. Cible les recettes à réévaluer."),
]

HOT_LISTS = [
    ("Heading2", "Listes chaudes (Hot Lists)"),
    ("Normal", "Comptage partiel : seuls les items choisis sont comptés ; les autres restent extrapolés."),
    ("Heading3", "Créer une Hot List"),
    ("Normal", "Count Inventory → Hot Lists → New → déplacez les items vers la droite → Save."),
    ("Heading3", "Compter avec une Hot List"),
    ("Normal", "New Inventory → Hot List → Finish. Plusieurs listes : cochez Multiple count sheets required."),
    ("Heading3", "Autres usages"),
    ("Normal", "Filtre Hot List dans Usage Summary ; type Hot List à l'impression des feuilles de comptage."),
]

REVALUING = [
    ("Heading2", "Réévaluation d'inventaire (Revaluing Inventory)"),
    ("Normal", "Si la valeur Usage Summary est incorrecte après correction de prix ou factures :"),
    ("Heading3", "Étapes"),
    ("Normal", "1. Corrigez prix, case size et conversions sur la fiche item."),
    ("Normal", "2. Vérifiez Period Purchases dans Usage Summary."),
    ("Normal", "3. Settings → Preferences → Inventory → Last Cost → Save."),
    ("Normal", "4. Ouvrez l'inventaire concerné ; modifiez le comptage (ex. 1 → 0 → 1) → Summarize → Save."),
    ("Normal", "5. Répétez sur inventaire de clôture si nécessaire ; rafraîchissez Usage Summary."),
]

PAR_LEVELS = [
    ("Heading2", "Niveaux Par des items (Item Par Levels)"),
    ("Normal", "Définissent un stock minimum (alerte) et maximum (cible après commande) par item."),
    ("Heading3", "Configurer"),
    ("Normal", "Inventory → Item Par Levels → Minimum Reorder Level et Maximum Reorder Level. Key Item et Actualize : voir chapitre 3."),
    ("Heading3", "Utilisation quotidienne"),
    ("Normal", "Tuile Order Reminder sur le dashboard ; case Include items below par level lors d'une commande ; combinable avec Forecasting Orders."),
    ("Heading3", "Rappels de commande (Order Reminders)"),
    ("Normal", "Par Levels configurés → tuile Order Reminder → Create Order → Round to Case → Save par fournisseur."),
]

ORDER_APPROVALS = [
    ("Heading2", "Approbation des commandes (Order Approvals)"),
    ("Normal", "Contrôle des coûts : saisie par tous, approbation requise avant export."),
    ("Heading3", "Activation"),
    ("Normal", "Settings → Preferences → Purchasing → Require order approval → Save."),
    ("Heading3", "Permissions"),
    ("Normal", "Security → Access Levels → Purchasing → Orders → Commands → Approve order (coché pour responsables, décoché pour gestionnaires)."),
    ("Heading3", "Flux"),
    ("Normal", "Manager crée la commande → responsable Set Approved → Export activé. Un manager peut exporter une commande déjà approuvée."),
]

SPLIT_CASE = [
    ("Heading3", "Commande en fraction de caisse (Split Case)"),
    ("Normal", "Cochez Split Case pour commander une fraction de caisse ; le système bascule vers l'unité secondaire."),
]

TAX_GROUPS = [
    ("Heading2", "Taxes, ajustements et groupes de taxes (Tax Groups)"),
    ("Normal", "Essentiel pour bière, alcool et items à prix TTC (taxes et dépôts inclus dans le prix affiché)."),
    ("Heading3", "Créer taxes et ajustements"),
    ("Normal", "Settings → Taxes and Adjustments → Add : Code, Recoverable, Percentage ou Value, Account."),
    ("Heading3", "Créer un Tax Group"),
    ("Normal", "Tax Groups → Add → cochez taxes applicables → Save."),
    ("Heading3", "Assigner à un item"),
    ("Normal", "Items → Case Size → Tax Group → Price including tax pour prix TTC. Recalculez les anciennes factures manuellement si besoin."),
    ("Heading3", "Préférence indispensable"),
    ("Normal", "Preferences → Purchasing → Show advanced tax adjustment fields in invoice window → Save."),
    ("Heading3", "Saisir une facture"),
    ("Normal", "Purchasing → Invoices → New : Supplier, Date, Invoice Number, lignes (double-clic catalogue), Taxes and Adjustments, Account Balance = 0 → Save."),
    ("Normal", "Onglet Taxes and Adjustments → vérifiez Account Balance avant Save."),
]

BEERSTORE = [
    ("Heading2", "Factures Beer Store / Brewers Retail (Canada)"),
    ("Normal", "Factures complexes — maîtrisez d'abord les Tax Groups."),
    ("Heading3", "Lecture de la facture"),
    ("Normal", "Order Number = n° facture ; Delivery Date = date ; totaux HST et dépôts sur plusieurs pages ; fuel surcharge et retours manuscrits."),
    ("Heading3", "Saisie"),
    ("Normal", "Brewers Retail → items → quantités → Expenses (fuel, retours) → Taxes/Adjustments avec ajustements manuels → total = dernière page."),
]

SYSCO_SFTP = [
    ("Heading3", "Migration Sysco : FTP vers SFTP"),
    ("Normal", "Invoice Import → Import Config → Sysco → sFTP et sB2B → Test credentials → Save."),
]

SALES_EXPORT = [
    ("Heading2", "Utilitaire d'export des ventes (Sales Export Utility)"),
    ("Normal", "Première méthode d'import POS si intégration configurée (raccourci bureau)."),
    ("Heading3", "Import quotidien"),
    ("Normal", "Double-clic → plage de dates → Process → attendez Success."),
    ("Heading3", "Mise à jour automatique"),
    ("Normal", "Dans l'utilitaire : Settings → Check for updates on startup → Save."),
    ("Heading3", "Récupérer des ventes manquantes"),
    ("Normal", "Relancez l'utilitaire sur la plage de dates manquante si des jours absents dans Sales."),
    ("Heading3", "Import CSV"),
    ("Normal", "Prérequis : Settings → Preferences → POS — dossier import + Pos Specification (engrenage) configurés. Puis Sales → New → … → fichier CSV → Continue → Save."),
]

ENTERPRISE_EXTRA = [
    ("Heading2", "Créer et déployer un Recipe Set"),
    ("Normal", "Enterprise → Recipe → Recipe Sets → New Recipe Set → sélection magasins et éléments → Save → Recipe Set Installer au magasin."),
    ("Heading2", "Entités en attente (Pending Entities)"),
    ("Normal", "Company → Preferences → require approval. Utility → Approved Pending Entities → Approve (propage) ou Ignore (reste au magasin). Non traité = erreurs dans les rapports."),
]

OC_MOBILE = [
    ("Heading1", "Chapitre 19 — OC Mobile : comptage sur appareil mobile"),
    ("Normal", "Module complémentaire iOS (iPhone/iPad) : export des feuilles de comptage, scan code-barres, réimport."),
    ("Heading2", "Configuration"),
    ("Normal", "Settings → Preferences → OC Mobile → identifiants → Test Login → Save."),
    ("Heading2", "Export et comptage"),
    ("Normal", "Count Inventory → Export → To OC Mobile. Sur l'appareil : Sign In → saisir quantités → Upload."),
    ("Heading2", "Import"),
    ("Normal", "Import → From OC Mobile → Merge (plusieurs appareils) ou Overwrite (un seul) → Save → Summarize."),
]

SECURITY = [
    ("Heading1", "Chapitre 20 — Sécurité : employés et niveaux d'accès"),
    ("Normal", "Un compte par personne qui saisit des données."),
    ("Heading2", "Utilisateur local"),
    ("Normal", "Security → Employees → New. Décochez Active pour désactiver sans supprimer."),
    ("Heading2", "Niveaux d'accès"),
    ("Normal", "Full Access, Update Not Create, Read Only, No Access — par fenêtre."),
    ("Heading2", "Utilisateurs Enterprise"),
    ("Normal", "Site web → Utility → Users → courriel (mot de passe temporaire par email)."),
    ("Heading2", "Journal d'activité"),
    ("Normal", "Rapport log des actions par employé."),
]

FAQ_EXTRA = [
    ("Heading3", "Les ventes POS n'apparaissent pas depuis plusieurs jours"),
    ("Normal", "Mettez à jour Sales Export Utility et relancez l'import pour la plage manquante."),
    ("Heading3", "La valeur d'un item dans Usage Summary semble incorrecte"),
    ("Normal", "Voir Réévaluation d'inventaire, chapitre 6."),
    ("Heading3", "Impossible de désactiver un item"),
    ("Normal", "Retirez-le des recettes qui l'utilisent d'abord."),
    ("Heading3", "Erreur « entity already exists » en Enterprise"),
    ("Normal", "Traitez Approved Pending Entities sur le site web."),
]

# Texte visible sous les titres dont le contenu original est surtout dans des tableaux Word
POS_PREFS = [
    ("Heading2", "Import POS — Preferences → POS (OC v5)"),
    ("Normal", "Dans Optimum Control v5, la configuration d'import des ventes CSV se trouve dans Settings → Preferences → POS. L'ancien menu « Setup → Configure POS » n'existe plus sous ce nom."),
    ("Normal", "• Default POS import folder — dossier par défaut ouvert quand vous faites Sales → New → … (pointez vers votre dossier de fichiers CSV)."),
    ("Normal", "• POS configuration — sélectionnez ou créez une Pos Specification ; cliquez l'engrenage à droite pour l'assistant Import Specification Settings."),
    ("Normal", "• Update selling price after import — Never si vous ne voulez pas écraser les prix menu à chaque import."),
    ("Normal", "• Default Category when Creating Product — catégorie par défaut si vous créez un product depuis Pending Sales."),
    ("Normal", "• Type of sales to view in reporting — Gross Sales ou Net Sales selon votre comptabilité."),
    ("Normal", "• Import Sales at Start — laissez décoché sauf import automatique au démarrage."),
] + POS_SPEC_WIZARD_BLOCKS[1:] + POS_SPEC_FIELD_BLOCKS[1:]

INVENTORY_PREFS = [
    ("Heading2", "Préférences Inventory (Settings → Preferences → Inventory)"),
    ("Normal", "Cet écran règle la valorisation du stock, le comportement des feuilles de comptage et les paramètres des prep sheets. Ne pas confondre avec Key Item et Actualize Usage Values : ces deux options se configurent sur chaque fiche Item (pas sur l'écran Preps en v5)."),
    ("Heading3", "Valorisation du stock (Value Inventory Using)"),
    ("Normal", "• FIFO — valorise au prix des premiers achats (recommandé par défaut)."),
    ("Normal", "• Last Cost — valorise au dernier prix de caisse (fiche item) ; utilisé aussi pour la réévaluation d'inventaire."),
    ("Normal", "• Weighted Average — moyenne pondérée des achats de la période."),
    ("Heading3", "Comptage"),
    ("Normal", "• Ask to Summarize on Close — rappel avant de fermer : ne pas oublier Summarize (recommandé : coché)."),
    ("Normal", "• Require Reason for Inventory Adjustments — oblige un motif si vous changez une quantité (recommandé : coché)."),
    ("Normal", "• Countsheet Column Display — All Columns affiche achat, caisse (split) et recette ; à fixer avant le premier inventaire."),
    ("Normal", "• Print countsheet groupings on separate pages — une page par regroupement à l'impression."),
    ("Heading3", "Prep sheets et seuils"),
    ("Normal", "• Default Prep Margin (ex. 10 %) — marge ajoutée aux quantités calculées sur les prep sheets."),
    ("Normal", "• Prep Amount Factor (ex. 100 %) — facteur sur les quantités de prep."),
    ("Normal", "• Last Prep Amount Calculation Range — plage de dates pour Prep Par Levels."),
    ("Normal", "• Warning / Critical Threshold — alertes si le stock dépasse le par (×2 jaune, ×3 rouge)."),
    ("Heading3", "Options avancées"),
    ("Normal", "• Inventory Interface — [None] sauf intégration bar (Freepour, Bluestreak)."),
    ("Normal", "• Disable perpetual amounts for non-counted items — si coché, les items non comptés ne recalculent pas leur qty théorique."),
    ("Normal", "• Set negative quantities to zero — force les quantités négatives à zéro après comptage."),
]

SETUP_CATEGORIES = [
    ("Normal", "Optimum Control propose trois écrans Setup distincts : Categories, Inventory Groups et Sales Groups. Ne les confondez pas — chacun a un rôle différent, mais ils partagent le même référentiel de catégories."),
    ("Normal", "Setup → Categories : c'est ici que vous créez les catégories. Pour chaque ligne, renseignez le nom (ex. Food, Beverage) et le champ Income Account (compte de type Income, ex. Sales Food). Les ventes de cette catégorie y seront rattachées à l'export comptable et dans Usage Summary."),
    ("Normal", "Le champ Sales Cat. / Sales Category que vous voyez sur Inventory Groups, Sales Groups et ailleurs n'est pas un écran séparé : c'est un menu déroulant qui affiche exactement les catégories créées dans Categories. Si vous ajoutez Food dans Categories, Food apparaît dans ce dropdown partout."),
    ("Heading3", "Groupes d'inventaire (Inventory Groups)"),
    ("Normal", "Setup → Inventory Groups (écran séparé). Chaque ligne comporte :"),
    ("Normal", "• Group desc. — nom du groupe (Proteins, Produce, Paper…), assigné à chaque Item et Prep."),
    ("Normal", "• Sales Cat. — menu déroulant : choisissez une catégorie créée dans Categories (Food, Beverage…). Obligatoire pour ventiler les coûts dans Usage Summary."),
    ("Normal", "• Account — compte CostOfSales où coder les achats du groupe (Food Cost, Beverage Cost, Paper / Supplies)."),
    ("Heading3", "Groupes de ventes (Sales Groups)"),
    ("Normal", "Setup → Sales Groups (écran séparé). Attribués aux Products (recettes vendues). Colonnes : Group desc. (nom du groupe, aligné sur le POS) + Sales Cat. (même dropdown que ci-dessus). Le Group desc. peut être plus fin que la catégorie (ex. groupe POS « Entrées » rattaché à la Sales Cat. Food)."),
]

SECTION_INTROS = {
    "Vocabulaire essentiel": [
        ("Normal", "Les termes ci-dessous reviennent partout dans Optimum Control. Le tableau qui suit les résume ; retenez surtout la chaîne Item → Prep → Product et la différence entre Sales Mix (ventes par produit) et Usage Summary (écart réel vs théorique)."),
    ],
    "Raccourcis clavier utiles": [
        ("Normal", "Raccourcis les plus utiles au quotidien : TAB (champ suivant), SHIFT + TAB (champ précédent), flèches haut/bas (navigation dans une liste), CTRL + C / X / V (copier, couper, coller). Le tableau ci-dessous reprend ces raccourcis."),
    ],
    "Volume et poids courants": [
        ("Normal", "Références de conversion les plus fréquentes en cuisine : 1 once liquide ≈ 29,5 ml ; 1 tasse (8 oz) = 236 ml ; 1 litre = 33,81 oz ; 1 once (poids) = 28,35 g ; 16 onces = 1 livre (454 g) ; 1 gallon US = 3,78 L. Le tableau ci-dessous détaille ces équivalences."),
    ],
    "Méthodes de tri des feuilles de comptage": [
        ("Normal", "Quatre modes de tri sont disponibles à l'impression ou à la saisie : Location (par emplacement de stockage, puis par groupe), Group (par groupe alphabétique), Category (par catégorie Food, Beer, etc.) et Custom (ordre shelf-to-sheet défini manuellement via Count Inventory → Customize Sort). Le tableau ci-dessous compare ces options."),
    ],
    "Saisir les comptages": [
        ("Normal", "Les quantités se saisissent sur la feuille de comptage (count sheet), y compris pour l'inventaire d'ouverture : OC n'exige aucune facture préalable — vous entrez ce qui est physiquement présent (ou les quantités de départ de votre scénario). Voir la section Inventaire d'ouverture ci-dessus si elle précède ce paragraphe."),
    ],
    "Les quatre états d'une vente importée": [
        ("Normal", "Chaque ligne importée du POS est classée dans l'un de ces états : Importée et liée (Product déjà associé — aucune action), Mismatched (description POS légèrement différente — Switch ou Unlink Product), Unlinked (aucun Product lié — Link To ou Create Product), Pending (mis de côté pour traitement ultérieur via Pending Sales Mix). Les identifiants volontairement exclus du calcul sont traités à part (section « Ignorer certains identifiants POS »)."),
    ],
    "Lier les produits POS aux Products Optimum Control": [
        ("Normal", POS_ID_NOTE),
        ("Normal", "La fonction Link Products to POS remplit Description, POS ID#, prix et Sales Group depuis l'export — il reste à ajouter les ingrédients. Si la recette existe déjà sans PLU, clic droit sur la ligne d'import → associer sans dupliquer."),
    ],
    "Saisir les ventes quotidiennes": [
        ("Normal", "La Daily Sales est la photographie financière du jour (Z de caisse), distincte du Sales Mix qui détaille les quantités vendues par product. Gross est souvent prérempli depuis le mix ; vous complétez remises, dépôts, main-d'œuvre et statistiques clients."),
    ],
    "Les achats de Joe": [
        ("Normal", "Joe achète chaque ingrédient dans le conditionnement fournisseur habituel. Le tableau ci-dessous indique, pour chaque item du hamburger, le coût de la caisse, la taille d'achat et l'unité de suivi."),
    ],
    "Calcul du coût unitaire de chaque ingrédient": [
        ("Normal", "Principe : diviser le coût total de la caisse par le nombre d'unités utilisables. Exemples — pains : 1,65 $ ÷ 12 = 13,7 cents par pain ; boulettes : 20,43 $ ÷ 66 ≈ 30,9 cents chacune ; ketchup : 34,00 $ ÷ 600 oz ≈ 5,6 cents l'once. Les calculs détaillés pour chaque ingrédient figurent dans les encadrés ci-dessous."),
    ],
    "Deux types de base de données": [
        ("Normal", "En Enterprise, deux rôles coexistent : Head Office (siège social — seul endroit autorisé pour créer ou modifier items, preps et products, et pour assigner les Consolidation ID) et Child (établissement — consultation en lecture seule, réception des mises à jour du siège). Le tableau ci-dessous résume ces droits."),
    ],
}

def flatten(src, dst):
    shutil.copy2(src, dst)
    with zipfile.ZipFile(dst) as z: f = {n: z.read(n) for n in z.namelist()}
    f["word/styles.xml"] = re.sub(r"<w:outlineLvl w:val=\"\d+\"/>", "", f["word/styles.xml"].decode()).encode()
    x = re.sub(r"<w:collapsed[^/]*/>", "", f["word/document.xml"].decode())
    x = re.sub(r"<w:sdt>\s*<w:sdtPr>.*?</w:sdtPr>\s*<w:sdtContent>(.*?)</w:sdtContent>\s*</w:sdt>", r"\1", x, flags=re.DOTALL)
    f["word/document.xml"] = x.encode()
    with zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as z:
        for n, d in f.items(): z.writestr(n, d)

def scan_empty(doc):
    n = 0
    for i, p in enumerate(doc.paragraphs):
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        if not t or not sn.startswith("Heading"): continue
        lv = hlevel(sn)
        has = False
        for q in doc.paragraphs[i+1:i+8]:
            qt = q.text.strip()
            qs = q.style.name if q.style else ""
            if not qt: continue
            if qs.startswith("Heading") and hlevel(qs) <= lv: break
            has = True; break
        if not has: n += 1
    return n

def fix_pos_legacy_text(doc):
    """Remplace l'ancien chemin Setup → Configure POS (obsolète en v5)."""
    old = "Setup → Configure POS"
    new = (
        "Settings → Preferences → POS "
        "(l'écran avec Default POS import folder et POS configuration — l'ancien libellé "
        "« Setup → Configure POS » n'existe plus en v5)"
    )
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
            n += 1
    if n:
        print("OK fix POS legacy text x", n)

def main():
    shutil.copy2(SRC, DOC)
    doc = Document(DOC)
    fix_pos_legacy_text(doc)

    for heading, blocks in SECTION_INTROS.items():
        insert_right_after_heading(doc, heading, blocks)

    racc = find_p(doc, "Raccourcis clavier utiles")
    if racc:
        idx = pidx(doc, racc)
        prev = doc.paragraphs[idx - 1]
        while idx > 1 and not prev.text.strip():
            idx -= 1
            prev = doc.paragraphs[idx - 1]
        insert_block_after(prev, DASHBOARD)
        print("OK dashboard apres navigation +", len(DASHBOARD))
    else:
        add_at_end_of_section(doc, "Navigation dans l'interface", DASHBOARD)
    add_at_end_of_section(doc, "Allergènes", DEACTIVATE + BARCODE, after="Chapitre 3 — Items")
    add_at_end_of_section(doc, "Les trois sections de la fiche Item",
        ITEM_CREATE_BLOCKS + item_multicasize_book_blocks() + TAX_GROUP_EXERCICE_BLOCKS[2:])
    add_at_end_of_section(doc, "Créer un Prep/Batch Recipe", PREPS_CREATE_BLOCKS + prep_book_blocks())
    add_at_end_of_section(doc, "Créer un Product (recette de menu)", PRODUCTS_CREATE_BLOCKS + product_book_blocks())
    add_at_end_of_section(doc, "Lier les produits POS aux Products Optimum Control", pos_id_book_blocks())
    add_at_end_of_section(doc, "Méthodes de tri des feuilles de comptage", opening_inventory_book_blocks())
    add_at_end_of_section(doc, "L'éditeur de fiche recette (Recipe Editor)", PREP_PAR + RECIPE_BOOK + MENU_REPORT)
    add_at_end_of_section(doc, "Supprimer un inventaire", HOT_LISTS + REVALUING)
    add_at_end_of_section(doc, "Utiliser les meilleures offres dans une commande", PAR_LEVELS + ORDER_APPROVALS + SPLIT_CASE)
    add_at_end_of_section(doc, "Conversions d'unités (Unit Conversions)", INVENTORY_PREFS)
    add_at_end_of_section(doc, "Catégories de vente (Sales Categories)", SETUP_CATEGORIES)
    add_at_end_of_section(doc, "Modifier la date d'une facture", TAX_GROUPS + BEERSTORE + SYSCO_SFTP + invoice_book_blocks())
    add_at_end_of_section(doc, "Configurer l'import POS", POS_PREFS + sales_till_tape_book_blocks())
    add_at_end_of_section(doc, "Saisir les ventes quotidiennes", daily_sales_book_blocks())
    add_at_end_of_section(doc, "Détailler les pertes (Waste Detailing)", SALES_EXPORT)
    add_at_end_of_section(doc, "Réception côté établissement", ENTERPRISE_EXTRA)
    add_at_end_of_section(doc, "Les montants de vente ne correspondent pas à ceux du point de vente", FAQ_EXTRA)
    add_before_conclusion(doc, OC_MOBILE + SECURITY)

    doc.save(DOC)
    flatten(DOC, PLAT)

    d2 = Document(DOC)
    empty = scan_empty(d2)
    print(f"\nTitres vides restants: {empty}")

    # verify Waste and Taxes
    for needle in ["Détailler les pertes", "Taxes, ajustements", "Modifier la date", "Réception côté"]:
        for i, p in enumerate(d2.paragraphs):
            if needle in p.text:
                nxt = d2.paragraphs[i+1].text.strip()[:50]
                print(f"  {needle[:30]} -> {nxt}")
                break

if __name__ == "__main__":
    main()
