#!/usr/bin/env python3
"""Génère le corrigé détaillé de l'exercice Comptoir du Marché."""
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from exercice_locale import publish_docx

OUT = Path(r"E:\OC DOCS\Exercice_Corrige_Detaille.docx")
CALC = Path(r"E:\OC DOCS\corrigé_calc.txt")

def add_title(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_p(doc, text):
    doc.add_paragraph(text)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

def build():
    subprocess.run(["python", r"E:\OC DOCS\calc_corrigé.py"], check=True)
    calc_text = CALC.read_text(encoding="utf-8") if CALC.exists() else ""

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CORRIGÉ DÉTAILLÉ\n")
    r.bold = True
    r.font.size = Pt(20)
    t.add_run("Exercice pratique — Le Comptoir du Marché\n\n").font.size = Pt(14)
    t.add_run("Ne consultez ce document qu'après avoir tenté l'exercice.\n").font.size = Pt(11)
    doc.add_page_break()

    add_title(doc, "Contexte local — Haïti / $US")
    add_p(doc, "Restaurant : Le Comptoir du Marché, Port-au-Prince, Haïti.")
    add_p(doc, "Devise : $US. Taxe : TCA 10 % (Taxe sur le Chiffre d'Affaires).")
    add_p(doc, "Fournisseurs : Distrib. Caraïbes, Boulangerie Pétion, Emballages Haïti, Boissons Nationale.")

    add_title(doc, "1. Coûts unitaires des items")
    add_p(doc, "Formule : coût caisse ÷ quantité en unité recette × (100 ÷ yield %).")
    add_table(doc, ["Item", "Calcul", "Coût unitaire"], [
        ["Pain burger", "2,40 $US ÷ 12", "0,20 $US / chacun"],
        ["Bœuf 80/20", "42,00 $US ÷ 5000 g", "0,0084 $US / g"],
        ["Bacon", "18,50 $US ÷ ~66 tranches (2×500 g)", "~0,295 $US / tranche"],
        ["Mayonnaise", "48,00 $US ÷ 15 120 ml", "0,00318 $US / ml"],
        ["Ketchup", "38,00 $US ÷ ~20 232 ml", "0,00188 $US / ml"],
        ["Pommes de terre", "22,00 $US ÷ (50 lb en g), yield 85 %", "0,00114 $US / g utilisable"],
        ["Huile friture", "38,00 $US ÷ 16 000 ml", "0,00238 $US / ml"],
        ["Boîte burger", "45,00 $US ÷ 500", "0,09 $US / chacun"],
        ["Serviette", "35,00 $US ÷ 5000", "0,007 $US / chacun"],
        ["Cola", "12,00 $US ÷ 24", "0,50 $US / canette"],
    ])

    add_title(doc, "2. Coûts des preps")
    add_table(doc, ["Prep", "Calcul", "Coût unité sortie"], [
        ["Boulette 80 g", "77 g bœuf + 3 g sel", "≈ 0,652 $US / boulette"],
        ["Frites maison — batch", "(5000 g pommes + 200 ml huile + 50 g sel) ÷ 4000 g rendu", "~0,00157 $US / g frite"],
        ["Portion frites 200 g", "200 g × coût/g batch", "~0,32 $US / portion"],
        ["Sauce burger", "(800 ml mayo + 200 ml ketchup) ÷ 950 ml", "~0,00307 $US / ml"],
        ["Chili — batch", "Haricots + bœuf chili + tomates + oignons + sel → 12 000 g", "~0,00309 $US / g"],
    ])

    add_title(doc, "3. Coûts recette et food cost %")
    add_p(doc, "Valeurs calculées (arrondis OC peuvent différer légèrement) :")
    add_table(doc, ["Product", "Coût recette", "Prix vente", "Food cost %", "Commentaire"], [
        ["Burger classique", "~1,13 $US", "14,95 $US", "~7,6 %", "Nourriture seule ; papier en sus"],
        ["Burger bacon", "~1,72 $US", "16,95 $US", "~10,1 %", "+ 2 tranches bacon"],
        ["Frites moyennes", "~0,32 $US", "5,50 $US", "~5,8 %", "1 Portion frites 200 g (prep) + serviette"],
        ["Salade César", "~1,43 $US", "13,50 $US", "~10,6 %", "4 items — sans poulet"],
        ["Bol chili", "~1,25 $US", "9,95 $US", "~12,5 %", "400 g chili + fourchette"],
        ["Cola", "0,50 $US", "3,25 $US", "15,4 %", "Item direct"],
        ["Eau", "0,33 $US", "2,75 $US", "12,1 %", "Item direct"],
    ])
    add_p(doc, "Note pédagogique : avec les prix d'achat de l'exercice, le food cost est nettement sous la cible de 32 % — vos marges sont confortables. Si OC affiche un % beaucoup plus élevé, vérifiez les unités (grammes vs kilos, ml vs litres) ou un yield oublié.")

    add_title(doc, "4. Détail — Burger classique")
    add_table(doc, ["Ingrédient", "Qté", "Coût unit.", "Coût ligne"], [
        ["Pain", "1", "0,20 $US", "0,20 $US"],
        ["Boulette (prep)", "1", "≈ 0,652 $US", "≈ 0,652 $US"],
        ["Sauce maison", "15 ml", "0,00307 $US/ml", "0,046 $US"],
        ["Tomates", "30 g", "0,0044 $US/g", "0,132 $US"],
        ["Boîte + serviette", "1+1", "—", "0,097 $US"],
        ["TOTAL", "", "", "~1,13 $US"],
    ])

    add_title(doc, "5. Ventes de la semaine — totaux")
    add_table(doc, ["POS ID", "Product", "Lun", "Mar", "Mer", "Jeu", "Ven", "TOTAL"], [
        ["101", "Burger classique", "25", "28", "22", "30", "35", "140"],
        ["102", "Burger bacon", "10", "12", "8", "15", "18", "63"],
        ["201", "Frites", "30", "35", "28", "40", "45", "178"],
        ["301", "Salade César", "8", "10", "12", "9", "14", "53"],
        ["401", "Bol chili", "15", "18", "20", "16", "22", "91"],
        ["501", "Cola", "40", "45", "38", "50", "55", "228"],
        ["502", "Eau", "12", "15", "10", "18", "20", "75"],
    ])
    add_p(doc, "Usage idéal approximatif bœuf burger (item Bœuf haché 80/20) : (140 + 63) boulettes × 77 g ≈ 15 631 g, plus chili et autres recettes.")

    add_title(doc, "6. Pertes — impact attendu")
    add_table(doc, ["Date", "Item", "Effet sur Usage Summary"], [
        ["Mardi", "Laitue 500 g", "Actual > Ideal sur laitue ; écart positif"],
        ["Mercredi", "Frites prep 400 g", "Actual prep frites augmente"],
        ["Jeudi", "Chili 800 g", "Perte prep ; shelf life expirée"],
        ["Vendredi", "Pain × 6", "Actual pain > ideal proportionnel"],
    ])

    add_title(doc, "7. Emplacements de stock — réponses")
    add_bullets_simple = [
        "Primary = emplacement de référence pour commandes et total consolidé.",
        "Secondary = où le stock est aussi physiquement (comptage par Location).",
        "Déplacer physiquement du stock entre armoires ne crée pas de transaction OC : seul le comptage par emplacement reflète la réalité.",
        "Tri Location à l'inventaire : feuille Réfrigérateur, puis Congélateur, etc.",
    ]
    for b in add_bullets_simple:
        doc.add_paragraph(b, style="List Bullet")

    add_title(doc, "8. Transferts entre établissements — flux")
    add_table(doc, ["Étape", "Magasin", "Menu", "Action"], [
        ["1", "Restaurant", "Stock Transfer → Request Stock", "Demande bœuf, pommes, cola à l'entrepôt"],
        ["2", "Entrepôt", "Approve Requests", "Approuve la demande"],
        ["3", "Entrepôt", "Fulfill Requisitions", "Expédie (partiel si stock insuffisant → backorder)"],
        ["Alt.", "Entrepôt", "Transfers → New", "Transfert direct sans réquisition"],
    ])
    add_p(doc, "Après fulfill : stock ↓ entrepôt, stock ↑ restaurant. Les deux magasins doivent partager les mêmes items (Item and Store Management).")

    add_title(doc, "9. Usage Summary — réponses types")
    for b in [
        "Plus grand écart $US : souvent bœuf haché ou pommes de terre (volume + ventes élevées).",
        "Ideal = 0, Actual > 0 : item sans recette liée ou Actualize activé (ex. huile si non liée aux frites).",
        "Food cost global : depend du mix ventes ; comparer catégorie Food seule.",
        "Drill-down Purchases : doit lister factures Distrib. Caraïbes, Boulangerie Pétion, etc.",
        "Drill-down Ideal : burger → products 101/102 ; frites → product 201.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_title(doc, "10. Défis bonus — solutions")
    add_title(doc, "Défi A", 2)
    add_p(doc, "Inventory → Adjust Inventory → ouvrir l'inventaire d'ouverture ou le dernier → corriger pommes de terre 2 cs → 1 cs → Save.")
    add_title(doc, "Défi B", 2)
    add_p(doc, "Modifier recette Burger classique (+10 g bacon ou 1 tranche) → Sales → Rerun Sales Mix sur la période.")
    add_title(doc, "Défi C", 2)
    add_p(doc, "Miscellaneous → Item Management → Eau → décocher Active → Save. Historique ventes conservé.")
    add_title(doc, "Défi D", 2)
    add_p(doc, "Corriger prix bœuf sur item → Preferences → Inventory → Last Cost → ajuster comptage inventaire (1→0→1) → Summarize → Usage Summary.")

    add_title(doc, "Annexe — Sortie brute du calculateur")
    for line in calc_text.splitlines():
        doc.add_paragraph(line)

    publish_docx(doc, OUT)

if __name__ == "__main__":
    build()
