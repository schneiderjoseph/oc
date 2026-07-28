from docx import Document
v = Document(r"E:\OC DOCS\Comprendre_Optimum_Control_V2.docx")
needles = ["Détailler les pertes", "Les quatre états", "Vocabulaire", "Raccourcis", "Méthodes de tri"]
lines = []
for needle in needles:
    for i, p in enumerate(v.paragraphs):
        if needle in p.text:
            lines.append(f"\n=== {needle} @ {i} ===")
            for j in range(i, min(i+5, len(v.paragraphs))):
                t = v.paragraphs[j].text.strip()[:70] if v.paragraphs[j].text.strip() else "(vide)"
                lines.append(f"  {j}|{t}")
            break
open(r"E:\OC DOCS\verify_v2.txt", "w", encoding="utf-8").write("\n".join(lines))
